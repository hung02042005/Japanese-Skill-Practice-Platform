"""Các hàm hỗ trợ python-docx dùng lại được, để build tài liệu RDS/SDS/Final-Release
từ các template FPT SWP391 trong Temp_Document/.

Sở dĩ cần các hàm này vì python-docx không có sẵn API kiểu "nhân bản section này N lần"
hay "thay ảnh mẫu trong template bằng ảnh thật". Mọi hàm ở đây đều thao tác trên một
docx.Document đã mở thật sự — không bao giờ tạo từ Document() rỗng — nhờ vậy font chữ,
style, header/footer và theme của template được giữ nguyên, không bị mất.

Import file này từ một script dùng một lần cho mỗi lần build tài liệu, ví dụ:

    import sys
    sys.path.insert(0, r".claude/skills/docx-report/scripts")
    from docx_tools import *

    doc = open_from_template("Temp_Document/Template3_SDS Document.docx",
                              "docs/07-Release-Documents/SDS_Document.docx")
    ... điền nội dung vào ...
    doc.save("docs/07-Release-Documents/SDS_Document.docx")
"""

import copy
import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


# --------------------------------------------------------------------------- #
# Mở file / lưu file
# --------------------------------------------------------------------------- #

def open_from_template(template_path, output_path):
    """Copy template sang output_path rồi mở BẢN COPY đó. Tuyệt đối không sửa
    trực tiếp file template gốc trong Temp_Document/."""
    template_path = Path(template_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(template_path, output_path)
    return Document(str(output_path))


# --------------------------------------------------------------------------- #
# Trích xuất text ở mức thấp (thuộc tính .text của paragraph không đáng tin cậy
# khi paragraph có nhiều run/drawing)
# --------------------------------------------------------------------------- #

def _element_text(element):
    """Ghép toàn bộ node <w:t> nằm dưới một element XML thô bất kỳ
    (paragraph hoặc table), theo đúng thứ tự xuất hiện trong tài liệu."""
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _has_drawing(element):
    return element.find(".//" + qn("w:drawing")) is not None


_CONTENT_TAGS = (qn("w:p"), qn("w:tbl"))


def _is_content_element(element):
    """Chỉ trả True cho paragraph/table thật sự (con trực tiếp của body). Word
    bọc field Mục Lục (Table of Contents) trong một content-control <w:sdt>,
    mà nội dung cache của nó chứa TOÀN BỘ tiêu đề (heading) có trong tài liệu —
    nếu tìm kiếm substring một cách ngây thơ trên mọi con của body, nó sẽ khớp
    vào bản cache của Mục Lục trước, chứ không phải paragraph heading thật.
    Luôn bỏ qua sdt/sectPr/bookmark/... khi so khớp."""
    return element.tag in _CONTENT_TAGS


# --------------------------------------------------------------------------- #
# Tìm kiếm
# --------------------------------------------------------------------------- #

def find_paragraphs_containing(doc, substr):
    """Tất cả paragraph cấp cao nhất của body có text chứa substr, trả về dạng Paragraph."""
    return [p for p in doc.paragraphs if substr in p.text]


def find_paragraph_index(doc, substr, occurrence=0):
    """Vị trí (index) trong list(doc.element.body.iterchildren()) của phần tử
    con thứ N (đếm từ 0) là paragraph-hoặc-table có text (đã ghép phẳng) chứa
    substr. Bỏ qua các con không phải nội dung thật (ví dụ field Mục Lục
    <w:sdt> — xem _is_content_element) để không bao giờ nhầm một mục trong
    Mục Lục với paragraph/heading thật mà nó trỏ tới."""
    children = list(doc.element.body.iterchildren())
    hits = [i for i, c in enumerate(children) if _is_content_element(c) and substr in _element_text(c)]
    if occurrence >= len(hits):
        raise ValueError(f"'{substr}' occurrence {occurrence} not found ({len(hits)} matches)")
    return hits[occurrence]


def find_image_paragraphs(doc):
    """Tất cả phần tử paragraph của body có chứa ảnh nhúng (<w:drawing>), theo
    đúng thứ tự trong tài liệu. Đây chính là các vị trí [IMAGE] (ảnh diagram
    mẫu) trong template RDS/SDS."""
    body = doc.element.body
    return [c for c in body.iterchildren() if c.tag == qn("w:p") and _has_drawing(c)]


def find_table_after(doc, heading_substr, occurrence=0):
    """<w:tbl> đầu tiên xuất hiện sau paragraph thứ N khớp với heading_substr."""
    children = list(doc.element.body.iterchildren())
    start = find_paragraph_index(doc, heading_substr, occurrence)
    for c in children[start + 1:]:
        if c.tag == qn("w:tbl"):
            return Table(c, doc)
    raise ValueError(f"No table found after '{heading_substr}'")


# --------------------------------------------------------------------------- #
# Sửa text của paragraph (các placeholder hướng dẫn kiểu [...] / <<...>>)
# --------------------------------------------------------------------------- #

def set_paragraph_text(paragraph, text):
    """Thay toàn bộ text của một Paragraph bằng `text`, giữ nguyên style của
    paragraph và định dạng của run đầu tiên (font/size/bold vẫn giữ nguyên)."""
    if isinstance(text, (list, tuple)):
        text = "\n".join(text)
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


def insert_paragraph_after(paragraph, text="", style=None):
    """docx không có sẵn API "chèn paragraph phía sau" — nhân bản một <w:p>
    rỗng, chèn nó vào bằng addnext, rồi bọc lại thành đối tượng Paragraph."""
    new_p = copy.deepcopy(paragraph._p)
    # loại bỏ mọi run/drawing bị copy theo từ paragraph nguồn
    for child in list(new_p):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def replace_image_in_paragraph(doc, paragraph_element, image_path, width_in=6.5):
    """Với một <w:p> thô đang chứa ảnh diagram mẫu (một <w:drawing>), xoá hết
    run của nó rồi chèn ảnh PNG thật vào thay thế, vẫn giữ nguyên
    alignment/style của paragraph. `doc` phải là Document (hoặc _Body) sở hữu
    paragraph đó để python-docx có thể tìm ra .part khi nhúng ảnh — nếu truyền
    None vào đây sẽ bị AttributeError ở sâu bên trong add_picture."""
    para = Paragraph(paragraph_element, doc)
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    para.add_run().add_picture(str(image_path), width=Inches(width_in))


# --------------------------------------------------------------------------- #
# Sửa cell / row của table
# --------------------------------------------------------------------------- #

def set_cell_text(cell, text):
    """Đặt text cho một cell của table, gộp về một paragraph duy nhất, vẫn
    giữ định dạng của run đầu tiên (nhờ vậy style header/body của template
    không bị mất)."""
    if isinstance(text, (list, tuple)):
        text = "\n".join(str(t) for t in text)
    paragraphs = cell.paragraphs
    set_paragraph_text(paragraphs[0], str(text))
    for extra in paragraphs[1:]:
        extra._p.getparent().remove(extra._p)


def add_table_row_like_last(table, values):
    """Nhân bản row cuối cùng của table (giữ nguyên style/border của cell) rồi
    điền `values` (list, mỗi phần tử ứng với một cột) vào bản nhân bản đó.
    Dùng hàm này cho các bảng tra cứu đơn giản cần nhiều row thật hơn số row
    mẫu (1-2 row) có sẵn trong template (bảng Actors, bảng Package, bảng Mô Tả
    Table trong DB, ...)."""
    last_tr = table.rows[-1]._tr
    new_tr = copy.deepcopy(last_tr)
    last_tr.addnext(new_tr)
    new_row = table.rows[-1]  # python-docx đọc lại rows từ XML mỗi lần truy cập
    for cell, value in zip(new_row.cells, values):
        set_cell_text(cell, value)
    return new_row


# --------------------------------------------------------------------------- #
# Nhân bản cả một section (heading + table + ảnh đi cùng nhau) — phần khó nhất
# --------------------------------------------------------------------------- #

def repeat_block(doc, start_substr, end_substr, n_copies, start_occurrence=0):
    """Deep-copy mọi phần tử con của body nằm giữa paragraph/table khớp
    start_substr và paragraph/table khớp end_substr (tính cả hai đầu), rồi
    chèn thêm n_copies - 1 bản sao ngay sau bản gốc.

    Trả về danh sách gồm n_copies "khoảng phần tử" (mỗi khoảng là list các
    lxml element thô), theo đúng thứ tự tài liệu: ranges[0] là khoảng GỐC (vẫn
    nằm nguyên vị trí cũ, text template vẫn còn nguyên — coi nó là mục #1 và
    vẫn phải điền nội dung cho nó), ranges[1:] là các bản sao mới cho mục
    #2..#N. Điền nội dung cho từng khoảng bằng fill_tokens_in_range() hoặc
    bằng cách tìm table/paragraph bên trong đúng khoảng đó (không bao giờ tìm
    kiếm lại trên toàn bộ tài liệu, nếu không sẽ sửa nhầm sang bản sao khác).
    """
    body = doc.element.body
    children = list(body.iterchildren())
    start_i = find_paragraph_index(doc, start_substr, start_occurrence)
    end_i = None
    for i in range(start_i, len(children)):
        if _is_content_element(children[i]) and end_substr in _element_text(children[i]):
            end_i = i
            break
    if end_i is None:
        raise ValueError(f"end marker '{end_substr}' not found after start marker")

    block = children[start_i:end_i + 1]
    ranges = [block]
    anchor = block[-1]
    for _ in range(n_copies - 1):
        clones = [copy.deepcopy(el) for el in block]
        for el in clones:
            anchor.addnext(el)
            anchor = el
        ranges.append(clones)
    return ranges


def fill_tokens_in_range(range_elements, mapping):
    """Thay mọi lần xuất hiện của từng token `{{TOKEN}}` (key trong mapping)
    bằng giá trị tương ứng, chỉ trong phạm vi các node <w:t> nằm bên trong
    range_elements đã cho."""
    for el in range_elements:
        for node in el.iter(qn("w:t")):
            if not node.text:
                continue
            text = node.text
            for token, value in mapping.items():
                if token in text:
                    text = text.replace(token, str(value))
            node.text = text


def tables_in_range(doc, range_elements):
    """Trả về các đối tượng Table (python-docx) cho mọi <w:tbl> nằm trong range_elements."""
    return [Table(el, doc) for el in range_elements if el.tag == qn("w:tbl")]


def image_paragraphs_in_range(range_elements):
    """Các <w:p> thô nằm trong range_elements mà có chứa ảnh diagram mẫu."""
    return [el for el in range_elements if el.tag == qn("w:p") and _has_drawing(el)]


# --------------------------------------------------------------------------- #
# Kiểm tra lại (bước 7 trong SKILL.md — môi trường này không có công cụ render
# docx headless, nên kiểm tra cấu trúc bằng code + xem lại bằng mắt trong Word
# là toàn bộ vòng QA)
# --------------------------------------------------------------------------- #

def summarize(doc):
    headings = {}
    for p in doc.paragraphs:
        if p.style and p.style.name.startswith("Heading"):
            headings[p.style.name] = headings.get(p.style.name, 0) + 1
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
        "headings": headings,
    }


def find_leftover_markers(doc):
    """Các paragraph/cell còn sót lại đoạn text hướng dẫn gốc của FPT (dạng
    ngoặc [...] hoặc <<...>>) — nếu còn sót nghĩa là có section bị bỏ sót
    trong lúc lắp ráp tài liệu."""
    leftovers = []
    for p in doc.paragraphs:
        t = p.text
        if ("[" in t and "]" in t) or ("<<" in t and ">>" in t):
            leftovers.append(t.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text
                if ("[" in t and "]" in t) or ("<<" in t and ">>" in t):
                    leftovers.append(t.strip())
    return leftovers
