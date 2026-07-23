# -*- coding: utf-8 -*-
import sys, copy
sys.path.insert(0, r".skill/docx-report/scripts")
from docx_tools import *
from docx.oxml.ns import qn

doc = Document("docs/07-Release-Documents/RDS_Document.docx")

def del_by_text(doc, substr, occurrence=0):
    children = list(doc.element.body.iterchildren())
    i = find_paragraph_index(doc, substr, occurrence)
    children[i].getparent().remove(children[i])

def del_range_before(doc, start_substr, end_substr):
    children = list(doc.element.body.iterchildren())
    i1 = find_paragraph_index(doc, start_substr)
    i2 = find_paragraph_index(doc, end_substr)
    for c in children[i1:i2]:
        c.getparent().remove(c)

N = 24
ranges = repeat_block(doc, "1.1 <<SubFeature Name>>", "<<transaction(s)>>", N)

SCREENS = [
dict(name="Login Screen", uc="UC-01",
 fields=[("Email","Text Box","placeholder=email@example.com"),
         ("Mật khẩu","Password Box","Có toggle hiện/ẩn mật khẩu"),
         ("Quên mật khẩu?","Hyperlink","href=/forgot-password"),
         ("Đăng nhập","Button","Submit -> POST /api/auth/login"),
         ("Tiếp tục với Google","Button","GoogleLogin component (@react-oauth/google), client-side ID token"),
         ("Đăng ký ngay","Hyperlink","href=/register")],
 db=[("student_users","R","Xác thực Email/Password; kiểm tra status/locked_until/login_attempts"),
     ("student_users","U","Cập nhật login_attempts, locked_until, last_login_at sau mỗi lần thử"),
     ("auth_tokens","C","Lưu refresh_token mới sau khi đăng nhập thành công")],
 sql="""SELECT id, email, password_hash, status, login_attempts, locked_until
FROM student_users WHERE email = ? AND is_deleted = 0;

UPDATE student_users SET login_attempts = 0, locked_until = NULL, last_login_at = GETDATE()
WHERE id = ?;

INSERT INTO auth_tokens (actor_type, student_id, token_type, token_value, expires_at, created_at)
VALUES ('STUDENT', ?, 'REFRESH', ?, DATEADD(DAY,7,GETDATE()), GETDATE());"""),

dict(name="Register Screen", uc="UC-02",
 fields=[("Họ và tên","Text Box","placeholder=Nguyễn Văn A"),
         ("Email","Text Box","type=email"),
         ("Mật khẩu","Password Box","Kèm thanh đo độ mạnh mật khẩu"),
         ("Xác nhận mật khẩu","Password Box","Validate khớp Mật khẩu"),
         ("Tạo tài khoản","Button","Submit -> POST /api/auth/register"),
         ("Đăng ký với Google","Button","GET /api/auth/oauth/google"),
         ("Gửi lại mã (cooldown 60s)","Button","POST /api/auth/resend-verification")],
 db=[("student_users","C","Tạo user mới, status='pending'"),
     ("student_users","R","Kiểm tra email đã tồn tại -> 409 EMAIL_EXISTS"),
     ("auth_tokens","C","Sinh mã OTP 6 số (token_type=EMAIL_VERIFICATION), hết hạn 10 phút")],
 sql="""SELECT id FROM student_users WHERE email = ? AND is_deleted = 0;

INSERT INTO student_users (full_name, email, password_hash, status, created_at)
VALUES (?, ?, ?, 'pending', GETDATE());

INSERT INTO auth_tokens (actor_type, student_id, token_type, token_value, expires_at)
VALUES ('STUDENT', ?, 'EMAIL_VERIFICATION', ?, DATEADD(MINUTE,10,GETDATE()));"""),

dict(name="Forgot Password Screen", uc="UC-03",
 fields=[("Email","Text Box","autoFocus, autocomplete=email"),
         ("Gửi link đặt lại mật khẩu","Button","Submit -> POST /api/auth/forgot-password"),
         ("Quay lại đăng nhập","Hyperlink","href=/login")],
 db=[("student_users","R","Tìm user theo email (không tiết lộ kết quả ra response)"),
     ("auth_tokens","C","Sinh token reset (>=32 bytes, hạn 15 phút) nếu user tồn tại")],
 sql="""SELECT id, email FROM student_users WHERE email = ? AND is_deleted = 0;

INSERT INTO auth_tokens (actor_type, student_id, token_type, token_value, expires_at)
VALUES ('STUDENT', ?, 'PASSWORD_RESET', ?, DATEADD(MINUTE,15,GETDATE()));"""),

dict(name="Reset Password Screen", uc="UC-03",
 fields=[("Mật khẩu mới","Password Box","Kèm thanh đo độ mạnh, autoFocus"),
         ("Xác nhận mật khẩu mới","Password Box","Validate khớp Mật khẩu mới"),
         ("Đặt lại mật khẩu","Button","Submit -> POST /api/auth/reset-password"),
         ("Gửi lại link mới","Hyperlink","href=/forgot-password, hiện khi token hết hạn/không hợp lệ")],
 db=[("auth_tokens","R","Kiểm tra token hợp lệ/chưa hết hạn/chưa dùng"),
     ("student_users","U","Cập nhật password_hash mới"),
     ("auth_tokens","U","Thu hồi (revoke) toàn bộ session + password_reset token khác")],
 sql="""SELECT student_id, expires_at, revoked_at FROM auth_tokens
WHERE token_value = ? AND token_type = 'PASSWORD_RESET';

UPDATE student_users SET password_hash = ? WHERE id = ?;

UPDATE auth_tokens SET revoked_at = GETDATE()
WHERE student_id = ? AND token_type IN ('REFRESH','PASSWORD_RESET') AND revoked_at IS NULL;"""),

dict(name="Kanji List Screen", uc="UC-07",
 fields=[("Level Tab (N5-N1)","Button (Tab)","Chọn cấp độ JLPT lọc lưới kanji"),
         ("Progress Bar","Progress Bar","% đã học / tổng kanji ở level đang chọn"),
         ("Kanji Cell","Button","Ô kanji trong lưới 8 cột, click -> /kanji/:id"),
         ("Pagination","Pagination Control","PAGE_SIZE=50")],
 db=[("kanji","R","Danh sách kanji theo level, phân trang, is_deleted=0"),
     ("student_content_progress","R","LEFT JOIN lấy isCompleted từng kanji")],
 sql="""SELECT k.id, k.character_value, k.meaning, k.jlpt_level,
       CASE WHEN p.status='completed' THEN 1 ELSE 0 END AS is_completed
FROM kanji k
LEFT JOIN student_content_progress p ON p.content_id=k.id AND p.content_type='kanji' AND p.student_id=?
WHERE k.jlpt_level=? AND k.status='published' AND k.is_deleted=0
ORDER BY k.id OFFSET ? ROWS FETCH NEXT 50 ROWS ONLY;"""),

dict(name="Kanji Practice (Writing/OCR) Screen", uc="UC-20",
 fields=[("Onyomi/Kunyomi/Nghĩa/Số nét","Text Display","Thông tin đọc âm/nghĩa/số nét kanji"),
         ("Stroke order image","Image Display","Ảnh tĩnh thứ tự nét viết"),
         ("Canvas vẽ","Canvas","Học viên vẽ lại ký tự bằng chuột/cảm ứng"),
         ("Nộp bài","Button","Export Canvas->PNG, gửi POST /api/submissions/handwriting"),
         ("Kết quả (Độ tương đồng %)","Progress Bar / Gauge","similarity_percent trả về sau khi AI OCR xử lý xong"),
         ("← Kanji trước / Kanji tiếp theo →","Button","Điều hướng sang kanji trước/sau")],
 db=[("kanji","R","Lấy chi tiết kanji theo id (stroke_order_url, onyomi, kunyomi, meaning)"),
     ("student_submissions","C","Tạo submission (handwriting, pending) khi nộp ảnh"),
     ("student_submissions","R","Poll trạng thái/kết quả (similarity_percent, is_correct, status=ai_graded)")],
 sql="""SELECT id, character_value, onyomi, kunyomi, meaning, stroke_count, stroke_order_url
FROM kanji WHERE id = ? AND is_deleted = 0;

INSERT INTO student_submissions (student_id, submission_type, kanji_id, handwriting_image_url, status, submitted_at)
VALUES (?, 'HANDWRITING', ?, ?, 'PENDING', GETDATE());

SELECT recognized_character, similarity_percent, is_correct, status
FROM student_submissions WHERE submission_id = ?;"""),

dict(name="Vocabulary List Screen", uc="UC-09",
 fields=[("Level Tab (N5-N1)","Button (Tab)","Chọn cấp độ JLPT"),
         ("Chủ đề (Topic)","Combo Box","Lọc từ vựng theo chủ đề"),
         ("Tìm từ vựng","Text Box (search)","Debounce 400ms"),
         ("Phát âm (▶)","Button","Phát audio, không autoplay"),
         ("+ FC (Thêm Flashcard)","Button","POST /api/flashcards"),
         ("✓ (Đánh dấu đã học)","Button","POST /api/learning-progress"),
         ("Pagination","Pagination Control","PAGE_SIZE=20")],
 db=[("vocabulary","R","Danh sách theo level/topic/search, phân trang"),
     ("student_content_progress","U","Đánh dấu từ đã học"),
     ("flashcards","C","Thêm từ vào flashcard cá nhân")],
 sql="""SELECT v.id, v.word, v.furigana, v.meaning, v.jlpt_level, v.audio_url
FROM vocabulary v
WHERE v.jlpt_level=? AND v.status='published' AND v.is_deleted=0 AND v.word LIKE '%'+?+'%'
ORDER BY v.id OFFSET ? ROWS FETCH NEXT 20 ROWS ONLY;

INSERT INTO flashcards (student_id, content_type, content_id, is_system, added_reason, next_review_date)
VALUES (?, 'VOCABULARY', ?, 0, 'learn', GETDATE());"""),

dict(name="Flashcard Session Screen", uc="Không nêu mã UC — SPEC riêng feat-flashcard-srs",
 fields=[("Progress Bar","Progress Bar","Tiến độ phiên (idx/queue.length)"),
         ("Thẻ NEW (Lật thẻ)","Card (flip component)","Thẻ từ mới, lật xem nghĩa/ví dụ/audio"),
         ("✗ Không nhớ / △ Khó / ✓ Dễ","Button (x3)","Tự đánh giá mức nhớ (chỉ thẻ NEW)"),
         ("Lựa chọn nghĩa (Quiz)","Button (option)","2-4 lựa chọn cho thẻ ÔN, server chấm đúng/sai"),
         ("Thêm vào Sổ tay","Button","Thêm từ trả lời sai vào Sổ tay cuối phiên")],
 db=[("flashcards","R","Dựng hàng đợi thẻ MỚI+ÔN theo deckId hoặc level+topic"),
     ("flashcards","U","Cập nhật SRS (ease_factor, interval_days, next_review_date) sau khi chấm"),
     ("flashcards","C","Thêm các từ trả lời sai vào Sổ tay (review deck)")],
 sql="""UPDATE flashcards
SET ease_factor=?, interval_days=?, repetition_count=?, next_review_date=?, last_reviewed_at=GETDATE(), last_rating=?
WHERE flashcard_id=? AND student_id=?;

SELECT * FROM flashcards
WHERE student_id=? AND last_session_id=? AND content_type='VOCABULARY' AND last_rating='WRONG';"""),

dict(name="Mock Test List Screen", uc="UC-10",
 fields=[("Level filter tabs (N5-N1)","Tab/Button group","Mặc định theo user.jlptLevel"),
         ("Số câu / Thời gian / Điểm đậu","Text (display)","totalQuestions, durationMin, passScore/maxScore"),
         ("Lần thi gần nhất","Text (display)","Ngày thi, điểm, đậu/không đậu; 'Chưa thi lần nào' nếu null"),
         ("Bắt đầu thi","Button","Điều hướng /mock-test/:id/attempt")],
 db=[("assessments","R","Danh sách đề thi type=exam theo level, phân trang"),
     ("test_attempts","R","Lấy thông tin lần thi gần nhất cho mỗi đề")],
 sql="""SELECT id, title, jlpt_level, duration_min, pass_score, total_score
FROM assessments WHERE assessment_type='exam' AND jlpt_level=? AND status='published';

SELECT TOP 1 * FROM test_attempts
WHERE student_id=? AND parent_id=? AND attempt_type='exam' ORDER BY started_at DESC;"""),

dict(name="Mock Test Attempt Screen", uc="UC-10",
 fields=[("ExamTopBar — Timer","Text (display, live)","mm:ss, đổi màu đỏ khi <5 phút"),
         ("ExamTopBar — Nộp bài","Button","Mở modal xác nhận nộp bài"),
         ("Đáp án A/B/C/D","Radio button group","Chọn 1 đáp án cho câu hỏi hiện tại"),
         ("Audio player","Audio control","Chỉ hiện khi câu hỏi có audioUrl (phần nghe)"),
         ("ExamNavigator — grid câu hỏi","Grid buttons","Nhảy nhanh tới câu, phân biệt đã trả lời/đang xem/chưa trả lời"),
         ("Modal xác nhận — Nộp bài xác nhận","Button","Submit thật sự")],
 db=[("assessments","R","Chi tiết đề thi gồm sections và questions"),
     ("questions","R","Câu hỏi, đáp án A/B/C/D, audio_url (KHÔNG trả correct_option)"),
     ("test_attempts","C","Tạo bản ghi khi bắt đầu (status=in_progress)"),
     ("attempt_answers","C","Lưu từng câu trả lời khi submit")],
 sql="""SELECT id, question_text, option_a, option_b, option_c, option_d, audio_url
FROM questions WHERE id IN (SELECT question_id FROM question_assignments WHERE parent_id=?)
ORDER BY display_order;

INSERT INTO test_attempts (student_id, attempt_type, parent_type, parent_id, started_at, status)
VALUES (?, 'EXAM', 'ASSESSMENT', ?, GETDATE(), 'IN_PROGRESS');"""),

dict(name="Mock Test Results Screen", uc="UC-10",
 fields=[("Điểm số (score/maxScore)","Text (display)","Tổng điểm đạt được / điểm tối đa"),
         ("Pass badge","Badge (display)","'ĐẬU'/'KHÔNG ĐẬU' + điểm cần đạt"),
         ("Section score bars","Progress bar (display)","% đạt theo từng kỹ năng (vocab/grammar/reading...)"),
         ("Bảng review từng câu","Table (display)","Câu, đáp án bạn chọn, đáp án đúng, đúng/sai"),
         ("Thi lại","Button","Điều hướng /mock-test/:id/attempt")],
 db=[("test_attempts","R","Kết quả lần thi (score, is_passed, sectionScores) theo attemptId"),
     ("attempt_answers","R","Chi tiết từng câu (selected_option, is_correct)"),
     ("assessments","R","assessmentTitle, jlptLevel, passScore liên kết với attempt")],
 sql="""SELECT score, max_score, is_passed, submitted_at FROM test_attempts WHERE attempt_id=? AND student_id=?;

SELECT aa.question_id, aa.selected_option, aa.is_correct, q.correct_option
FROM attempt_answers aa JOIN questions q ON q.id=aa.question_id WHERE aa.attempt_id=?;"""),

dict(name="Quiz & Luyện Tập Screen", uc="UC-11",
 fields=[("Level filter tabs","Tab/Button group","N5-N1, mặc định theo user.jlptLevel"),
         ("Kỹ năng (skill select)","Select dropdown","Tất cả/Từ vựng/Ngữ pháp/Đọc hiểu/Nghe"),
         ("QuizCard","Text (display)","title, skill, level, questionCount, bestScore"),
         ("Đáp án A/B/C/D","Radio button group","role=radio"),
         ("Nộp bài (x/y)","Button","Disable nếu chưa trả lời hết"),
         ("Kết quả — sao đánh giá","Text (display)","1-5 sao dựa trên scorePct")],
 db=[("assessments","R","Danh sách quiz theo level/skill, phân trang"),
     ("questions","R","Câu hỏi của 1 quiz (KHÔNG trả correct_option trước khi nộp)"),
     ("test_attempts","C","Tạo bản ghi khi submit"),
     ("attempt_answers","C","Lưu đáp án đã chọn + kết quả đúng/sai")],
 sql="""SELECT id, title, jlpt_level, question_count FROM assessments
WHERE assessment_type='quiz' AND jlpt_level=? AND status='published';

INSERT INTO attempt_answers (attempt_id, question_id, selected_option, is_correct)
VALUES (?, ?, ?, ?);"""),

dict(name="Dashboard Screen", uc="UC-19",
 fields=[("StreakCard — số ngày streak","Text (display)","VD: 42"),
         ("HeroBanner — Progress bar","Progress bar","{completed}/{total} bài"),
         ("HeroBanner — CTA 'HỌC TỪ MỚI'","Button","Điều hướng /learn/new"),
         ("LessonCard","Text/Icon (display)","Title, description, mini progress bar, JLPT badge"),
         ("QuickActionCard x4","Button/Card","Ôn Flashcard, Thi Thử JLPT, Từ Điển, Tiến Độ"),
         ("StatCard 'words'/'days'","Card (display)","Số từ đã học, số ngày học trong tháng")],
 db=[("student_users","R","Thông tin học viên đang đăng nhập, streak"),
     ("courses","R","Course hiện tại (title, jlptLevel, completedLessons, totalLessons)"),
     ("lessons","R","Danh sách bài học kèm trạng thái/progress"),
     ("student_content_progress","R","wordCount, daysThisMonth")],
 sql="Không đủ căn cứ trong tài liệu nguồn để viết SQL cụ thể — file SPEC-dashboard.md chỉ mô tả JSON API contract."),

dict(name="Speaking Practice Screen", uc="UC-13",
 fields=[("Level Tabs","Tab group","N5-N1, role=tablist"),
         ("SpeakingCard — targetText","Text (display)","Câu mẫu cần đọc"),
         ("▶ Nghe mẫu","Button","Phát audio mẫu (sampleAudioUrl)"),
         ("🎙 Bắt đầu ghi âm / ⏹ Dừng","Button","MediaRecorder API, timer 0:00/0:30 (MAX_RECORD_SECS)"),
         ("Nộp bài →","Button","Submit audio (handleSubmit)"),
         ("Kết quả — score/transcript/wordResults","Text (display)","Điểm %, bản phiên âm AI, từng từ đúng/sai")],
 db=[("lessons","R","Danh sách bài luyện speaking theo level"),
     ("student_submissions","C","Tạo lượt nộp audio, sinh submissionId"),
     ("student_submissions","R","Poll trạng thái/kết quả AI (score, transcript, wordResults)")],
 sql="Không đủ căn cứ trong tài liệu nguồn để viết SQL cụ thể — file SPEC-speaking.md chỉ mô tả JSON API contract."),

dict(name="Notebook Screen (Sổ Tay — Từ cần ôn lại)", uc="Không nêu mã UC — ref FR-FC-81",
 fields=[("Tab 'Tất cả' / 'Đến hạn'","Tab","Lọc theo isDue, kèm count"),
         ("Ô tìm kiếm 'Tìm trong sổ...'","Search input","Tìm client-side theo frontText"),
         ("▶ Ôn lại ngay (N)","Button","Điều hướng /review?deckId={reviewDeckId}"),
         ("NotebookWordCard — nút gỡ (🗑)","Button (icon)","Mở modal xác nhận gỡ khỏi sổ")],
 db=[("flashcard_decks (deck hệ thống 'Từ cần ôn lại')","R","Tìm deck isSystem=true"),
     ("flashcards","R","Liệt kê từ trong sổ theo deckId"),
     ("flashcards","D (soft-delete)","Gỡ 1 từ khỏi sổ — file ghi rõ ĐÂY LÀ GIẢ ĐỊNH cần xác nhận contract backend")],
 sql="Không đủ căn cứ trong tài liệu nguồn để viết SQL cụ thể — file SPEC-notebook.md chỉ mô tả JSON API contract."),

dict(name="Verify Email Screen", uc="Không nêu mã UC — ref UC-02-register.md",
 fields=[("Email","Input (text)","Prefill từ query param ?email=, cho sửa tay"),
         ("Mã xác minh (OTP)","Input (6 chữ số)","Nhập mã OTP nhận qua email"),
         ("Xác minh","Button","Submit -> POST /api/auth/verify-email"),
         ("Gửi lại mã xác minh","Button","POST /api/auth/resend-verification, cooldown hiển thị 'Gửi lại mã (Ns)'")],
 db=[("student_users","U","Cập nhật status=active sau khi xác minh thành công"),
     ("auth_tokens","R, U","Kiểm tra & vô hiệu hoá mã OTP sau khi dùng/hết hạn/sai quá 5 lần"),
     ("auth_tokens","C","Sinh mã OTP mới khi resend, giới hạn 1 lần/60 giây")],
 sql="""SELECT student_id, token_value, expires_at FROM auth_tokens
WHERE student_id=(SELECT id FROM student_users WHERE email=?) AND token_type='EMAIL_VERIFICATION';

UPDATE student_users SET status='active' WHERE id=?;"""),

dict(name="Staff Dashboard Screen", uc="Không nêu mã UC cụ thể",
 fields=[("Stat card — Nháp / Chờ duyệt","Stat card (clickable)","Click -> /staff/content?status=draft|pending_review"),
         ("Stat card — Tickets mở / Bài nói chờ chấm","Stat card (clickable)","Click -> /staff/tickets, /staff/grading"),
         ("Quick action — 4 thẻ","Card/link","Soạn Nội Dung, Ngân Hàng Câu Hỏi, Hỗ Trợ, Chấm Bài Nói"),
         ("Recent Activity table","Table (5 dòng)","Ngày, Loại, Tiêu đề, Trạng thái")],
 db=[("Bảng nội dung học liệu (course/lesson/vocab/grammar/kanji/question)","R","Đếm draftCount, pendingReviewCount"),
     ("tickets","R","Đếm openTicketCount"),
     ("student_submissions","R","Đếm pendingGradingCount (speaking chờ chấm)")],
 sql="Không đủ căn cứ trong tài liệu nguồn để viết SQL cụ thể."),

dict(name="Staff Content Management Screen", uc="UC-25, UC-27",
 fields=[("Search","Text input","Debounce 400ms"),
         ("Level filter / Status filter","Select","N5-N1 / draft-pending_review-published-rejected"),
         ("Content Type Tabs","Tab","Khóa học/Bài học/Từ vựng/Ngữ pháp/Kanji"),
         ("Content table","Table","Danh sách theo tab, cột Level/Trạng thái/Ngày cập nhật"),
         ("Nút Sửa/Gửi duyệt/Xem","Button","Theo điều kiện status"),
         ("Modal tạo/sửa — title/jlptLevel + field riêng theo loại","Text/Select/Textarea/Upload","Field khác nhau theo Khóa học/Bài học/Từ vựng/Ngữ pháp/Kanji"),
         ("Modal footer — Lưu nháp / Lưu và gửi duyệt","Button","status=draft hoặc tạo rồi gọi submit-review")],
 db=[("courses, lessons, vocabulary, grammar_points, kanji","C, R, U","Qua GET/POST/PUT /api/staff/{loại}"),
     ("Bảng nội dung tương ứng (status column)","U","Qua POST /api/staff/contents/submit-review -> pending_review")],
 sql="Không đủ căn cứ trong tài liệu nguồn để viết SQL cụ thể."),

dict(name="Staff Questions Screen", uc="UC-24",
 fields=[("Search / Kỹ năng filter / Level filter / Loại câu filter / Trạng thái filter","Text input/Select","Tìm & lọc câu hỏi"),
         ("Question table","Table","#, Câu hỏi (truncate 80 ký tự), Kỹ năng, Level, Loại, Trạng thái (+icon 🔒 nếu khoá)"),
         ("+ Thêm câu hỏi","Button","Mở QuestionFormModal"),
         ("Nút Tạo phiên bản mới","Button (hiện nếu isLocked)","Prefill form từ câu hỏi gốc, tạo câu hỏi mới độc lập"),
         ("Modal — Câu hỏi/Loại/Kỹ năng/Cấp độ (bắt buộc)","Textarea/Radio/Select","Nội dung + phân loại câu hỏi"),
         ("Modal — Đáp án A/B/C/D + Đáp án đúng","Text input/Radio","Chỉ hiện nếu loại=Trắc nghiệm")],
 db=[("questions (bao gồm đáp án)","C, R, U","Qua GET/POST/PUT /api/staff/questions"),
     ("questions (status column)","U","Qua POST /api/staff/contents/submit-review, contentType=question")],
 sql="Không đủ căn cứ trong tài liệu nguồn để viết SQL cụ thể."),

dict(name="Staff Grading Screen", uc="UC-31",
 fields=[("Status tabs","Tab","'Chờ chấm (pendingCount)' / 'Đã chấm'"),
         ("Submission List","List (card)","Tên, level, ngày, thời lượng, điểm AI/trạng thái"),
         ("Audio Player","<audio controls>","Nghe bài nộp (recordingUrl)"),
         ("AI Score Section","Display + progress bars","Điểm tổng, Phát âm, Lưu loát, lỗi phát hiện, gợi ý"),
         ("Điểm thủ công (manualScore)","Number input (0-100)","Nếu rỗng thì dùng điểm AI"),
         ("Phản hồi chi tiết (feedback)","Textarea (bắt buộc)","Nhận xét cho học viên"),
         ("Lưu điểm","Button","Gọi gradeSubmission()")],
 db=[("student_submissions (speaking)","R","GET /api/staff/submissions?submissionType=speaking — danh sách + chi tiết"),
     ("student_submissions (speaking)","U","POST /{id}/grade — cập nhật manual_score, manual_feedback, status=graded")],
 sql="""UPDATE student_submissions
SET manual_score=?, manual_feedback=?, status='GRADED', graded_by=?, graded_at=GETDATE()
WHERE submission_id=?;"""),

dict(name="Staff Review Queue Screen (Manager)", uc="UC-33, UC-34",
 fields=[("View Tabs","Tab button","'Chờ duyệt (n)' / 'Đã xuất bản'"),
         ("Type/Level/Submitted-by filter","Dropdown","Lọc theo contentType/jlptLevel/người gửi"),
         ("Review Queue Table","Table","Loại, Tiêu đề, Level, Người tạo, Gửi lúc"),
         ("Nút Duyệt / Từ chối / Yêu cầu sửa","Button","Mở ReviewActionModal"),
         ("ReviewActionModal — Lý do phản hồi (*)","Textarea (bắt buộc khi Reject/Request Changes)","Lý do"),
         ("Published Content Table — Nút Lưu trữ/Thu hồi","Button","Archive/Unpublish content")],
 db=[("Bảng nội dung (question/lesson/course/grammar/vocabulary/kanji/assessment)","R","GET /manager/review-queue, /manager/contents?status=published"),
     ("Bảng nội dung + review_log","U/I","POST /manager/reviews, /manager/reviews/request-changes"),
     ("Bảng nội dung","U","PUT /manager/contents/{id}/status (published/archived/unpublished)")],
 sql="Không đủ căn cứ trong tài liệu nguồn để viết SQL cụ thể."),

dict(name="Admin Dashboard Screen", uc="UC-36",
 fields=[("Stat card 'Tổng người dùng'","Stat value","Tổng student+staff+admin"),
         ("Stat card 'Hoạt động hôm nay'","Stat value","Distinct users đăng nhập 24h qua"),
         ("Stat card 'Trạng thái'","Stat value/badge","Hệ thống OK/BẢO TRÌ/LỖI"),
         ("Activity Log list","List (10 items)","Log hành động gần nhất"),
         ("Quick Action — 5 link","Link","Quản lý người dùng, SMTP, Bảo mật, Thông báo, Bảo trì")],
 db=[("student_users, staff_users, admin_users","R","Tổng người dùng, hoạt động hôm nay"),
     ("test_attempts","R","Số quiz attempts hôm nay"),
     ("admin_audit_logs","R","10 log mới nhất (logId, adminEmail, actionType, targetType, description, createdAt) — trực tiếp từ response mẫu trong file")],
 sql="Không đủ căn cứ trong tài liệu nguồn để viết SQL cụ thể, trừ admin_audit_logs đã có schema xác nhận trực tiếp trong file."),

dict(name="Admin Users Screen (ManageUsers)", uc="UC-37",
 fields=[("Reset Requests Panel — Cấp mật khẩu tạm","Button","Endpoint nêu trực tiếp trong file: POST /api/admin/staff/{id}/temp-password"),
         ("Type Tab — Học viên/Nhân viên/Quản trị","Tab","Chuyển activeType"),
         ("Search box / Trạng thái filter / Cấp độ JLPT filter / Vai trò filter","Input/Dropdown","Tìm & lọc user"),
         ("Table — Người dùng/Trạng thái/Cấp độ/Streak/Vai trò/Đăng nhập lần cuối","Avatar+Text/Badge","Danh sách user"),
         ("Action — Đình chỉ/Kích hoạt/Đặt lại mật khẩu/Đổi vai trò/Xóa mềm","Icon button","Thao tác quản trị"),
         ("Modal Suspend — Lý do đình chỉ","Textarea (10-500 ký tự, bắt buộc)","Lý do"),
         ("Modal Tạo Staff — Họ tên/Email/Vai trò","Input/Radio","Tạo Staff mới")],
 db=[("staff_users","U","POST /api/admin/staff/{id}/temp-password — cấp mật khẩu tạm (endpoint xác nhận trực tiếp)"),
     ("student_users, staff_users, admin_users","R","Danh sách/lọc user — endpoint không nêu rõ trong file, suy luận từ UI"),
     ("student_users, staff_users","U","Đình chỉ/Kích hoạt/Đổi vai trò/Xóa mềm — không có endpoint cụ thể trong file, suy luận từ hành vi UI")],
 sql="Không đủ căn cứ trong tài liệu nguồn để viết SQL cụ thể, ngoại trừ endpoint temp-password đã xác nhận."),

dict(name="Admin Settings Screen", uc="UC-39 (UC-40 Notification Rules đã gỡ khỏi FE, chỉ tham khảo lịch sử)",
 fields=[("Tab Hệ thống/Email-SMTP/Bảo mật","Tab","?tab=system|email|security"),
         ("Toggle 'Chế độ bảo trì'","Toggle switch","Bật/tắt maintenance mode, có confirm dialog"),
         ("smtp_host/port/secure/username/password/from_name","Input text/number/select/password","Cấu hình SMTP"),
         ("Kiểm tra kết nối","Button","Test SMTP — POST /api/admin/settings/smtp/test"),
         ("max_login_attempts / lockout_duration_minutes / jwt_expiry_minutes","Inline edit field (integer)","Cấu hình bảo mật")],
 db=[("system_settings","R","GET /api/admin/settings[/{group}]"),
     ("system_settings","U","PUT /api/admin/settings/{group}/{key}, bao gồm maintenance_mode")],
 sql="""SELECT setting_key, setting_value FROM system_settings WHERE setting_group=?;

UPDATE system_settings SET setting_value=? WHERE setting_group=? AND setting_key=?;"""),
]

assert len(SCREENS) == N, f"expected {N}, got {len(SCREENS)}"

for i, (rng, sc) in enumerate(zip(ranges, SCREENS), start=1):
    tabs = tables_in_range(doc, rng)
    ui_t, db_t = tabs[0], tabs[1]
    imgs = image_paragraphs_in_range(rng)

    # headings: [0]="1.i <<SubFeature Name>>", [1]="a. <<Screen/Function Name>>"
    ps_in_range = [el for el in rng if el.tag == qn("w:p")]
    sub_head = Paragraph(ps_in_range[0], doc)
    set_paragraph_text(sub_head, f"1.{i} {sc['name']}")
    scr_head = Paragraph(ps_in_range[1], doc)
    set_paragraph_text(scr_head, f"a. {sc['name']} (Liên quan: {sc['uc']})")

    # xoa cau huong dan [Provide brief description...] ngay sau scr_head neu con
    for p_el in ps_in_range[2:4]:
        p = Paragraph(p_el, doc)
        if p.text.startswith("[") or p.text == "<<Mockup prototype>>":
            set_paragraph_text(p, "")

    # anh mockup: xoa (chua co screenshot/mockup that), thay bang ghi chu TODO
    for img_el in imgs:
        replace_para = Paragraph(img_el, doc)
        for run in list(replace_para.runs):
            run._element.getparent().remove(run._element)
        replace_para.add_run("[Mockup/Screenshot: TODO — chưa có, xem SPEC-*.md gốc trong docs/03-Interface-Specs/feature-specs/frontend/]")

    # UI Design table: xoa dong merged "Field Group Name" (index 1) truoc khi dien
    fg_row = ui_t.rows[1]._tr
    fg_row.getparent().remove(fg_row)
    r0 = sc["fields"][0]
    set_cell_text(ui_t.rows[1].cells[0], r0[0]); set_cell_text(ui_t.rows[1].cells[1], r0[1]); set_cell_text(ui_t.rows[1].cells[2], r0[2])
    for r in sc["fields"][1:]:
        add_table_row_like_last(ui_t, list(r))

    # Database Access table: xoa dong ".." placeholder (index 2) truoc khi dien them
    if len(db_t.rows) > 2:
        ph_row = db_t.rows[2]._tr
        ph_row.getparent().remove(ph_row)
    r0d = sc["db"][0]
    set_cell_text(db_t.rows[1].cells[0], r0d[0]); set_cell_text(db_t.rows[1].cells[1], r0d[1]); set_cell_text(db_t.rows[1].cells[2], r0d[2])
    for r in sc["db"][1:]:
        add_table_row_like_last(db_t, list(r))

    # SQL Commands paragraph: tim trong range (sau db_t)
    sql_p = None
    for p_el in ps_in_range:
        p = Paragraph(p_el, doc)
        if p.text.startswith("[Provide the detailed SQL"):
            sql_p = p
            break
    if sql_p is not None:
        set_paragraph_text(sql_p, sc["sql"])

doc.save("docs/07-Release-Documents/RDS_Document.docx")
print("Section III (before cleanup GAMS examples) done.")
