"""Builds docs/07-Release-Documents/SDS_Document.docx from Temp_Document/Template3_SDS
Document.docx, following .skill/docx-report/SKILL.md. Run once:
    python docs/07-Release-Documents/build_sds.py
"""
import copy
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent.parent / ".skill/docx-report/scripts"))
from docx_tools import (
    open_from_template, set_paragraph_text, set_cell_text, add_table_row_like_last,
    repeat_block, tables_in_range, image_paragraphs_in_range, replace_image_in_paragraph,
    summarize, find_leftover_markers, _element_text, _is_content_element,
)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).parent.parent.parent
OUT = ROOT / "docs/07-Release-Documents/SDS_Document_v2.docx"
DIA_RDS = ROOT / "docs/07-Release-Documents/diagrams/rds"
DIA_SDS = ROOT / "docs/07-Release-Documents/diagrams/sds"

doc = open_from_template(str(ROOT / "Temp_Document/Template3_SDS Document.docx"), str(OUT))
body = doc.element.body


def find_in_range(range_elements, substr):
    for el in range_elements:
        if substr in _element_text(el):
            return el
    raise ValueError(f"'{substr}' not found in range")


def remove_range_slice(range_list, start_el, end_el_exclusive):
    """Remove elements from range_list (and from the XML tree) starting at
    start_el up to (not including) end_el_exclusive."""
    si = range_list.index(start_el)
    ei = range_list.index(end_el_exclusive)
    for el in range_list[si:ei]:
        el.getparent().remove(el)
    del range_list[si:ei]


def set_run_font(run, name):
    """Stamp ascii/hAnsi/cs/eastAsia all at once so Word never falls back to a
    theme font for any part of the run (the gap python-docx's Font.name setter
    leaves, since it only touches ascii/hAnsi)."""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)


def set_lines(paragraph, lines, mono=False):
    """Multi-line text in one paragraph using real <w:br/> line breaks."""
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    for i, line in enumerate(lines):
        if i > 0:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        if mono:
            set_run_font(run, "Consolas")


def resize_table_rows(table: Table, n_data_rows):
    while len(table.rows) - 1 < n_data_rows:
        add_table_row_like_last(table, [""] * len(table.columns))
    while len(table.rows) - 1 > n_data_rows:
        last = table.rows[-1]._tr
        last.getparent().remove(last)


def fill_method_table(table, methods):
    """methods: list of (no, method, description)"""
    resize_table_rows(table, len(methods))
    for i, (no, method, desc) in enumerate(methods):
        row = table.rows[i + 1]
        set_cell_text(row.cells[0], no)
        set_cell_text(row.cells[1], method)
        set_cell_text(row.cells[2], desc)


# --------------------------------------------------------------------------- #
# Cover page
# --------------------------------------------------------------------------- #

project_name_p = next(p for p in doc.paragraphs if "<<Project name>>" in p.text)
set_paragraph_text(project_name_p, "Japanese Skill Practice Platform")
for p in doc.paragraphs:
    if "Hanoi, August 2022" in p.text:
        set_paragraph_text(p, "– Hanoi, July 2026 –")
        break

# --------------------------------------------------------------------------- #
# Record of changes
# --------------------------------------------------------------------------- #

changes_table = doc.tables[0]
CHANGES = [
    ("2026-07-22", "A", "AI Agent (per user request)",
     "Initial SDS_Document.docx built from Template3: filled in the Overview (14 packages, DB schema) "
     "and 6 Code Designs flows (Auth, Quiz Submission, Flashcard SRS, Kanji Writing, Speaking Grading, Content Review)."),
    ("2026-07-22", "M", "AI Agent (per user request)",
     "Corrected the Speaking Submission Grading flow: added the Student submission step "
     "(package speaking - SpeakingController/SpeakingService) and updated SupportTicketService.manualGrade "
     "to match the current code (no longer requires status=AI_GRADED, only blocks REJECTED) "
     "— reflects the removal of AI auto-grading (commit db0d1648)."),
    ("2026-07-22", "M", "AI Agent (per user request)",
     "Translated the whole document to English and hardened the font handling for the SQL code blocks "
     "(explicit ascii/hAnsi/cs/eastAsia font stamp instead of relying on theme fallback)."),
]
resize_table_rows(changes_table, len(CHANGES))
for i, (date, amd, incharge, desc) in enumerate(CHANGES):
    row = changes_table.rows[i + 1]
    set_cell_text(row.cells[0], date)
    set_cell_text(row.cells[1], amd)
    set_cell_text(row.cells[2], incharge)
    set_cell_text(row.cells[3], desc)

# --------------------------------------------------------------------------- #
# I.1 Code Packages
# --------------------------------------------------------------------------- #

for p in doc.paragraphs:
    if "[Provide the package diagram" in p.text:
        set_paragraph_text(p, "")
        break

# find the first image paragraph after "1. Code Packages" heading, before "Package descriptions"
all_children = [c for c in body.iterchildren() if _is_content_element(c)]
start_idx = next(i for i, c in enumerate(all_children) if "1. Code Packages" in _element_text(c))
end_idx = next(i for i, c in enumerate(all_children) if "Package descriptions" in _element_text(c))
pkg_image_el = next(c for c in all_children[start_idx:end_idx] if c.tag == qn("w:p") and c.find(".//" + qn("w:drawing")) is not None)
replace_image_in_paragraph(doc, pkg_image_el, DIA_RDS / "package-diagram.png")

PACKAGES = [
    ("01", "auth", "Student login/register, JWT access + refresh token, email verification, forgot password, Google OAuth login."),
    ("02", "staff", "Staff/StaffManager accounts: separate login, staff member management, staff password reset."),
    ("03", "admin", "Admin accounts, audit log, statistics dashboard, system configuration (system_settings), maintenance mode."),
    ("04", "student", "Student profile, dashboard, learning progress (student_content_progress), avatar, enrolled courses."),
    ("05", "learning", "Learning content: Kana, Kanji, Vocabulary, Grammar, Lesson by JLPT level (N5-N1)."),
    ("06", "assessment", "Question bank, Quiz/Exam (Assessment), attempts (TestAttempt), server-side grading, speaking/writing submissions (StudentSubmission)."),
    ("07", "speaking", "Student records and submits speaking (shadowing) exercises, polls the result (SpeakingController/SpeakingService); grading is done by teachers in the support package — no AI auto-grading step anymore (removed by a recent refactor)."),
    ("08", "staffcontent", "Where Staff draft content before publishing: questions, quizzes, exams, grammar, vocabulary, and a dedicated Staff dashboard."),
    ("09", "contentreview", "Content review flow for Staff-authored content (approve/reject) before it goes public, with review audit logging."),
    ("10", "publishedcontent", "Snapshot of published content, served to Students (kept separate from Staff's draft copy)."),
    ("11", "flashcard", "Flashcard + Spaced Repetition (SRS), system decks or user-created decks."),
    ("12", "dictionary", "Kanji/Vocabulary dictionary lookup by keyword or type."),
    ("13", "notification", "Notifications sent to individual Students (notification type, delivery channel)."),
    ("14", "support", "Support tickets submitted by Students and staff replies (TicketReply); Staff's speaking-grading queue (SupportTicketService, UC-31)."),
]
pkg_table = doc.tables[1]
resize_table_rows(pkg_table, len(PACKAGES))
for i, (no, name, desc) in enumerate(PACKAGES):
    row = pkg_table.rows[i + 1]
    set_cell_text(row.cells[0], no)
    set_cell_text(row.cells[1], name)
    set_cell_text(row.cells[2], desc)

# --------------------------------------------------------------------------- #
# I.2 Database Design
# --------------------------------------------------------------------------- #

for p in doc.paragraphs:
    if "[Provide the tables relationship" in p.text:
        set_paragraph_text(p, "DBMS: Microsoft SQL Server 2019+, database JLPT_LearningDB. See "
                               "docs/02-SDD-Architecture/database-design/JLPT_database.md for the full 23-table "
                               "reference — the ER diagram below summarizes only the table groups relevant to the "
                               "6 flows detailed in Section II.")
        break

all_children = [c for c in body.iterchildren() if _is_content_element(c)]
start_idx = next(i for i, c in enumerate(all_children) if "a. Database Schema" in _element_text(c))
end_idx = next(i for i, c in enumerate(all_children) if "b. Table Description" in _element_text(c))
er_image_el = next(c for c in all_children[start_idx:end_idx] if c.tag == qn("w:p") and c.find(".//" + qn("w:drawing")) is not None)
replace_image_in_paragraph(doc, er_image_el, DIA_RDS / "er-diagram.png")

TABLES = [
    ("01", "student_users", "Student accounts, including OAuth (Google), current/target JLPT level, learning streak."),
    ("02", "staff_users", "Staff/StaffManager accounts."),
    ("03", "admin_users", "Admin accounts."),
    ("04", "auth_tokens", "Shared token store (session, refresh, reset password, verify email) for all 3 actor types via actor_type + actor_id."),
    ("05", "assessments", "JLPT Quiz and Exam (combined into a single table via assessment_type)."),
    ("06", "questions", "Question bank, with A/B/C/D answers stored inline."),
    ("07", "question_assignments", "Assigns a question to an assessment or a lesson via parent_type/parent_id."),
    ("08", "test_attempts", "One quiz/exam/practice attempt by a Student; keeps status to prevent double submission."),
    ("09", "attempt_answers", "A Student's answer to each question in an attempt, used for grading and review."),
    ("10", "flashcards", "One flashcard belonging to a Student (vocab/kanji/grammar/custom) plus SRS state (ease_factor, interval_days, next_review_date)."),
    ("11", "kanji_writing_attempts", "One Student attempt at hand-writing a Kanji character on canvas; average DTW score plus stroke_details (JSON)."),
    ("12", "student_submissions", "Submissions requiring Staff grading: submission_type = SPEAKING or HANDWRITING; keeps both the legacy AI score columns (ai_*) and the Staff score (manual_score, manual_feedback, graded_by)."),
    ("13", "admin_audit_logs", "Records every content-review / submission-grading action taken by Staff."),
]
db_table = doc.tables[2]
resize_table_rows(db_table, len(TABLES))
for i, (no, name, desc) in enumerate(TABLES):
    row = db_table.rows[i + 1]
    set_cell_text(row.cells[0], no)
    set_cell_text(row.cells[1], name)
    set_cell_text(row.cells[2], desc)

print("Overview section done.")

# --------------------------------------------------------------------------- #
# II. Code Designs — 6 features, repeat_block over the template's Feature1 block
# --------------------------------------------------------------------------- #

FEATURES = [
    {
        "name": "Authentication & Login",
        "class_diagram": DIA_SDS / "class-auth.png",
        "sequence_diagram": DIA_SDS / "seq-auth.png",
        "classes": [
            ("AuthenticationService", [
                ("01", "checkAccountType(email, ip)", "Input: email, caller IP. Rate-limited to 10 calls/minute/IP. Returns the account type (Student/Staff/Admin) or not-found so the FE can show the right login form."),
                ("02", "login(LoginRequest, ip)", "Authenticates a Student via AuthenticationManager + passwordEncoder; on success, issues an access token + refresh token (JwtProvider) and stores the refresh token in auth_tokens with actor_type=STUDENT."),
                ("03", "loginStaff(LoginRequest, ip)", "Same as login but for StaffUser/AdminUser; additionally checks the account status (locked/inactive)."),
                ("04", "refresh(RefreshTokenRequest)", "Input: refresh token. Looks up auth_tokens by hash, checks expires_at and that it has not been revoked, issues a new access token (rotating the refresh token)."),
                ("05", "logout(LogoutRequest)", "Revokes (marks) the current refresh token in auth_tokens — no hard-delete (ADR-004)."),
                ("06", "loginWithGoogle(GoogleTokenRequest)", "Verifies the Google ID token, finds or creates a student_users row by oauth_provider_id, issues a JWT the same way as a normal login."),
            ]),
            ("RegistrationService", [
                ("01", "register(RegisterRequest)", "Input: email, password, name. Hashes the password with bcrypt cost >= 10 (ADR-003), creates student_users with email_verified_at=null, and publishes a SendVerificationEmailEvent asynchronously."),
            ]),
        ],
        "sql": [
            "-- checkAccountType / login: look up the account by email",
            "SELECT * FROM student_users WHERE email = ? AND is_deleted = 0;",
            "",
            "-- store the refresh token on successful login",
            "INSERT INTO auth_tokens (actor_type, actor_id, token_type, token_hash, expires_at, created_at)",
            "VALUES ('STUDENT', ?, 'REFRESH', ?, ?, GETDATE());",
            "",
            "-- refresh(): look up a still-valid refresh token",
            "SELECT * FROM auth_tokens",
            "WHERE token_hash = ? AND token_type = 'REFRESH' AND expires_at > GETDATE();",
            "",
            "-- logout(): revoke the token (soft, no DELETE)",
            "UPDATE auth_tokens SET revoked_at = GETDATE() WHERE token_hash = ?;",
        ],
    },
    {
        "name": "Quiz Submission (Assessment)",
        "class_diagram": DIA_SDS / "class-quiz.png",
        "sequence_diagram": DIA_SDS / "seq-quiz.png",
        "classes": [
            ("QuizService", [
                ("01", "startQuiz(quizId, student)", "Input: quiz id, the current Student. Checks the quiz is PUBLISHED, loads the assigned questions (question_assignments), creates a new TestAttempt with status=IN_PROGRESS. Never returns the correct answers to the client."),
                ("02", "submitQuiz(quizId, studentId, attemptId, answers)", "Locks the test_attempts row with findByIdForUpdate (pessimistic lock) to prevent a race on duplicate submission; checks attempt.student.id==studentId (blocks submitting on someone else's behalf) and status==IN_PROGRESS (blocks double submission, LESSON-005). Calls calculateScore."),
                ("03", "calculateScore(attempt, assessment, answers)", "The score is computed entirely server-side, never trusting a score from the client: matches answers against each Question's correct_option, saves each AttemptAnswer, updates attempt.status=SUBMITTED."),
            ]),
        ],
        "sql": [
            "-- startQuiz(): check the quiz is published",
            "SELECT * FROM assessments",
            "WHERE id = ? AND assessment_type = 'QUIZ' AND status = 'PUBLISHED';",
            "",
            "-- load the assigned questions, in display order",
            "SELECT qa.* FROM question_assignments qa",
            "WHERE qa.parent_type = 'ASSESSMENT' AND qa.parent_id = ?",
            "ORDER BY qa.display_order;",
            "",
            "-- create a new attempt when the student starts the quiz",
            "INSERT INTO test_attempts (student_id, attempt_type, parent_type, parent_id, started_at, max_score, status)",
            "VALUES (?, 'QUIZ', 'ASSESSMENT', ?, GETDATE(), ?, 'IN_PROGRESS');",
            "",
            "-- submitQuiz(): lock the attempt row to prevent duplicate submission",
            "SELECT * FROM test_attempts WITH (UPDLOCK, ROWLOCK) WHERE id = ?;",
            "",
            "-- save each answer after grading",
            "INSERT INTO attempt_answers (attempt_id, question_id, selected_option, is_correct)",
            "VALUES (?, ?, ?, ?);",
            "",
            "-- update the attempt after submission",
            "UPDATE test_attempts",
            "SET status = 'SUBMITTED', submitted_at = GETDATE(), score = ?",
            "WHERE id = ?;",
        ],
    },
    {
        "name": "Flashcard SRS (Spaced Repetition)",
        "class_diagram": DIA_SDS / "class-flashcard.png",
        "sequence_diagram": DIA_SDS / "seq-flashcard.png",
        "classes": [
            ("FlashcardSrsService", [
                ("01", "submitReview(flashcardId, studentId, ReviewRequest)", "Verifies card ownership (ownCardOrThrow). For a multiple-choice vocabulary quiz, the server matches selectedOptionId against the correct contentId itself (never trusts a correct/incorrect flag from the client). Calls applySm2 to update the review schedule."),
                ("02", "getSession(studentId, deckId, topicId, newLimit)", "Input: exactly one of deckId (personal notebook) or topicId (curriculum topic), plus a limit on new cards (default 10, cap 20). Locks on (studentId, deck/topic) to prevent two tabs creating a session at the same time. Priority order: not-yet-learned -> due for review -> the rest."),
                ("03", "applySm2(Flashcard, LastRating)", "Implements the SM-2 algorithm: WRONG -> ease decreases (floor 1.30), repetitionCount resets; HARD -> ease unchanged; EASY -> ease+0.1 (cap 2.50). Always updates nextReviewDate = today + intervalDays."),
            ]),
        ],
        "sql": [
            "-- ownCardOrThrow(): verify card ownership before grading",
            "SELECT * FROM flashcards",
            "WHERE flashcard_id = ? AND student_id = ? AND is_deleted = 0;",
            "",
            "-- applySm2() -> save(): update the review schedule after grading",
            "UPDATE flashcards",
            "SET ease_factor = ?, interval_days = ?, repetition_count = ?,",
            "    next_review_date = ?, last_reviewed_at = GETDATE(), last_rating = ?",
            "WHERE flashcard_id = ?;",
            "",
            "-- end of session: collect words answered wrong in THIS session",
            "SELECT * FROM flashcards",
            "WHERE student_id = ? AND last_session_id = ? AND content_type = 'VOCABULARY' AND last_rating = 'WRONG';",
        ],
    },
    {
        "name": "Kanji Writing Evaluation (OCR/DTW Similarity)",
        "class_diagram": DIA_SDS / "class-kanji.png",
        "sequence_diagram": DIA_SDS / "seq-kanji.png",
        "classes": [
            ("KanjiWritingServiceImpl", [
                ("01", "evaluateStroke(EvaluateRequest)", "Input: the coordinates of the stroke the Student just drew + the reference stroke for that specific stroke index. Normalizes both paths (normalize+downsample to at most 20 points), computes the distance via DTW, maps it to a quality bucket using 3 fixed thresholds (300/650/1200). Synchronous, does not persist to the DB."),
                ("02", "saveAttempt(AttemptRequest, studentId)", "Input: every stroke drawn for one Kanji character + the per-stroke DTW score. Computes avgDtwScore as the average of the per-stroke scores, derives finalQuality, saves one KanjiWritingAttempt row (with strokeDetails as JSON)."),
                ("03", "computeDtw(s1, s2)", "Classic Dynamic Time Warping (ADR-007: compares similarity % only, no stroke-order analysis like image-based OCR). Invariant to differences in drawing speed between the two paths."),
            ]),
        ],
        "sql": [
            "-- saveAttempt(): persist one completed writing attempt",
            "INSERT INTO kanji_writing_attempts",
            "    (student_id, kanji_id, character_value, total_strokes, avg_dtw_score, final_quality, stroke_details, created_by)",
            "VALUES (?, ?, ?, ?, ?, ?, ?, ?);",
            "",
            "-- evaluateStroke() never queries the DB — all DTW computation happens in memory per request.",
        ],
    },
    {
        "name": "Speaking Submission Grading (Shadowing)",
        "class_diagram": DIA_SDS / "class-speaking.png",
        "sequence_diagram": DIA_SDS / "seq-speaking.png",
        "classes": [
            ("SpeakingService", [
                ("01", "getExercises(level, studentId)", "Input: JLPT level + student id. Returns the list of published speaking exercises (Lesson of type SPEAKING) for that level, with attemptCount/bestScore (the student's highest teacher-graded score)."),
                ("02", "submit(exerciseId, audio, student)", "Input: exercise id + multipart audio file. Validates the exercise exists & is PUBLISHED, stores the file via SpeakingAudioStorageService (no BLOB in the DB, LESSON-002), creates a StudentSubmission with status=PENDING — there is no AI grading step (removed by a refactor, see the Sequence Diagram). Returns a jobId for the Student to poll."),
                ("03", "getResult(jobId, studentId)", "Input: submission id + the calling student id (only ever returns that student's own submission). Returns PENDING if not yet graded (manualScore==null), COMPLETED with score/feedback once graded, FAILED if status=REJECTED."),
            ]),
            ("SupportTicketService", [
                ("01", "getAllSubmissions(submissionType, status, page, size)", "Input: submissionType filter (default speaking), status, pagination. Returns a Page<SubmissionResponse> for the Staff grading queue (UC-31)."),
                ("02", "getSubmissionDetail(submissionId)", "Returns the detail of one submission: recordingUrl, and if already graded, manualScore/manualFeedback/gradedBy. finalScore = manualScore if present, otherwise falls back to aiOverallScore (a legacy column, always null for new data since the AI grading step was removed)."),
                ("03", "manualGrade(submissionId, actorEmail, ManualGradeRequest)", "Current business rule: can only grade a submission whose submissionType=SPEAKING (422 otherwise) and status != REJECTED (422 if already rejected) — no longer requires status=AI_GRADED as it used to, since the AI grading step was removed (commit db0d1648). Sets manualScore/manualFeedback/gradedBy/gradedAt, moves status to GRADED, sends a Notification (ACHIEVEMENT) and writes an AdminAuditLog."),
            ]),
        ],
        "sql": [
            "-- SpeakingService.submit(): create a new submission, awaiting teacher grading (no AI step anymore)",
            "INSERT INTO student_submissions (student_id, submission_type, status, exercise_id, recording_url, created_at)",
            "VALUES (?, 'SPEAKING', 'PENDING', ?, ?, GETDATE());",
            "",
            "-- SupportTicketService.getAllSubmissions(): the speaking grading queue",
            "SELECT * FROM student_submissions",
            "WHERE submission_type = 'SPEAKING' AND (status = ? OR ? IS NULL)",
            "ORDER BY submitted_at ASC;",
            "",
            "-- manualGrade(): record the teacher's score, change status (no longer checks status=AI_GRADED)",
            "UPDATE student_submissions",
            "SET manual_score = ?, manual_feedback = ?, graded_by = ?, graded_at = GETDATE(), status = 'GRADED'",
            "WHERE submission_id = ? AND status != 'REJECTED';",
            "",
            "-- notify the Student once grading is done",
            "INSERT INTO notifications (student_id, title, message, notification_type, reference_key, created_by)",
            "VALUES (?, 'Your speaking submission has been graded', ?, 'ACHIEVEMENT', ?, ?);",
        ],
    },
    {
        "name": "Content Review (Manager Approve/Reject)",
        "class_diagram": DIA_SDS / "class-review.png",
        "sequence_diagram": DIA_SDS / "seq-review.png",
        "classes": [
            ("ContentReviewService", [
                ("01", "getContentDetail(managerEmail, contentId, typeStr)", "Verifies managerEmail is an active STAFF_MANAGER (requireManager), resolves the right handler for the type (course/lesson/grammar/vocabulary/kanji/question/assessment), returns a snapshot (title, JLPT level, author, submission date, content)."),
                ("02", "review(managerEmail, ReviewActionRequest)", "Blocks reviewing one's own content (guardSelfReview). APPROVE -> handler.approve() (status=published); REJECT -> feedback is required, then handler.transitionFromPending(REJECTED). Every action is logged via ReviewAuditService.log(...)."),
                ("03", "requestChanges(managerEmail, RequestChangesRequest)", "Asks Staff to revise content: feedback is required, moves status back to draft (default) or rejected depending on targetStatus. Uses the same guardSelfReview + audit-log mechanism as review()."),
                ("04", "ensureUpdated(affectedRows)", "If the UPDATE affected 0 rows, the content is no longer in pending_review (already handled concurrently by another Manager) -> throws ConcurrentReviewException to prevent one review overwriting another."),
            ]),
        ],
        "sql": [
            "-- findActiveById(): fetch content awaiting review (example: the GRAMMAR handler)",
            "SELECT * FROM grammar_points WHERE grammar_id = ? AND status = 'pending_review';",
            "",
            "-- approve(): only updates if still pending_review (prevents double approval)",
            "UPDATE grammar_points",
            "SET status = 'published', approved_by = ?, published_at = GETDATE()",
            "WHERE grammar_id = ? AND status = 'pending_review';",
            "",
            "-- reject() / requestChanges(): move back to draft or rejected",
            "UPDATE grammar_points",
            "SET status = ?, updated_at = GETDATE()",
            "WHERE grammar_id = ? AND status = 'pending_review';",
            "",
            "-- write an audit row for every review action",
            "INSERT INTO admin_audit_logs (staff_id, action, target_table, target_id, note, created_at)",
            "VALUES (?, ?, ?, ?, ?, GETDATE());",
        ],
    },
]

ranges = repeat_block(doc, "<Feature/Function Name1>", "[Provide the detailed SQL", n_copies=len(FEATURES))

for i, (feat, rng) in enumerate(zip(FEATURES, ranges), start=1):
    # heading2: "N. <name>"
    heading_el = find_in_range(rng, "<Feature/Function Name1>")
    set_paragraph_text(Paragraph(heading_el, doc), f"{i}. {feat['name']}")

    # drop instructional paragraph right under the heading
    instr = find_in_range(rng, "[Provide the detailed design")
    instr.getparent().remove(instr)
    rng.remove(instr)

    # a. Class Diagram
    instr = find_in_range(rng, "[This part presents the class diagram")
    instr.getparent().remove(instr)
    rng.remove(instr)
    imgs = image_paragraphs_in_range(rng)
    replace_image_in_paragraph(doc, imgs[0], feat["class_diagram"])

    # b. Class Specifications
    instr = find_in_range(rng, "[Provide the description for each class")
    instr.getparent().remove(instr)
    rng.remove(instr)

    xyz_heading = find_in_range(rng, "XYZ Class")
    xyz_instr = find_in_range(rng, " [Provide the detailed description")
    xyz_instr.getparent().remove(xyz_instr)
    rng.remove(xyz_instr)
    set_paragraph_text(Paragraph(xyz_heading, doc), f"{feat['classes'][0][0]} Class")
    xyz_table = tables_in_range(doc, rng)[0]
    fill_method_table(xyz_table, feat["classes"][0][1])

    if len(feat["classes"]) == 2:
        abc_heading = find_in_range(rng, "ABC Class")
        stray = find_in_range(rng, "Class Methods")
        stray.getparent().remove(stray)
        rng.remove(stray)
        abc_instr = find_in_range(rng, "[Provide the detailed description")
        abc_instr.getparent().remove(abc_instr)
        rng.remove(abc_instr)
        set_paragraph_text(Paragraph(abc_heading, doc), f"{feat['classes'][1][0]} Class")
        abc_table = tables_in_range(doc, rng)[1]
        fill_method_table(abc_table, feat["classes"][1][1])
    else:
        abc_heading = find_in_range(rng, "ABC Class")
        seq_heading = find_in_range(rng, "c. Sequence Diagram(s)")
        remove_range_slice(rng, abc_heading, seq_heading)

    # c. Sequence Diagram(s)
    instr = find_in_range(rng, "[Provide the sequence diagram(s)")
    instr.getparent().remove(instr)
    rng.remove(instr)
    imgs = image_paragraphs_in_range(rng)
    replace_image_in_paragraph(doc, imgs[0], feat["sequence_diagram"])

    # d. Database Queries
    sql_para_el = find_in_range(rng, "[Provide the detailed SQL")
    sql_para = Paragraph(sql_para_el, doc)
    set_lines(sql_para, feat["sql"], mono=True)

    print(f"Feature {i} ({feat['name']}) filled.")

# The template's trailing "2. <Feature/Function Name2>" stub (single angle
# brackets, not "<<...>>") sits right after our repeat_block's end marker, so
# it survives untouched — remove it and the stray glyph paragraph after it.
for p in list(doc.paragraphs):
    if "<Feature/Function Name2>" in p.text:
        nxt = p._p.getnext()
        if nxt is not None and nxt.tag == qn("w:p") and len(_element_text(nxt).strip()) <= 2:
            nxt.getparent().remove(nxt)
        p._p.getparent().remove(p._p)
        break

doc.save(str(OUT))
print("Saved", OUT)

summary = summarize(doc)
print(summary)
leftovers = find_leftover_markers(doc)
print(f"{len(leftovers)} leftover markers:")
for l in leftovers:
    print(" -", l[:120])
