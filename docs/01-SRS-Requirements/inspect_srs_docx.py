from docx import Document
from docx.oxml.ns import qn

path = r"D:\project\Japanese-Skill-Practice-Platform\Guides  Templates-20260721\Template1_SRS Document.docx"
doc = Document(path)

out = (
    r"C:\Users\admin\AppData\Local\Temp\claude\d--project-Japanese-Skill-Practice-Platform"
    r"\d70ebc54-2d65-490a-93f2-2d096b5a6372\scratchpad\srs_docx_structure.txt"
)


def has_drawing(el):
    return el.find(".//" + qn("w:drawing")) is not None


with open(out, "w", encoding="utf-8") as f:
    f.write(f"total paragraphs: {len(doc.paragraphs)}\n")
    f.write(f"total tables: {len(doc.tables)}\n")
    f.write(f"total inline images: {len(doc.inline_shapes)}\n\n")

    body = doc.element.body
    children = list(body.iterchildren())
    f.write(f"total body children: {len(children)}\n\n")

    table_idx = 0
    for i, el in enumerate(children):
        if el.tag == qn("w:p"):
            text = "".join(node.text or "" for node in el.iter(qn("w:t")))
            is_img = has_drawing(el)
            # find matching Paragraph style
            style = None
            for p in doc.paragraphs:
                if p._p is el:
                    style = p.style.name if p.style else None
                    break
            marker = " [IMAGE]" if is_img else ""
            if text.strip() or is_img:
                f.write(f"[{i}] P style={style!r}{marker}: {text[:120]!r}\n")
        elif el.tag == qn("w:tbl"):
            tbl = doc.tables[table_idx]
            table_idx += 1
            nrows = len(tbl.rows)
            ncols = len(tbl.columns)
            header = [c.text[:30] for c in tbl.rows[0].cells] if nrows else []
            f.write(f"[{i}] TABLE #{table_idx-1} {nrows}x{ncols} header={header}\n")
        elif el.tag == qn("w:sectPr"):
            f.write(f"[{i}] sectPr\n")
        else:
            f.write(f"[{i}] OTHER tag={el.tag}\n")

print("done", out)
