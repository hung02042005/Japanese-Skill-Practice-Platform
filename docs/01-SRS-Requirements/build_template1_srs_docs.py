from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "01-SRS-Requirements"
TEMPLATE = ROOT / "Guides  Templates-20260721" / "Template1_SRS Document.docx"


actors_vi = [
    ("1", "Guest", "Người dùng chưa đăng nhập: có thể đăng ký tài khoản, đăng nhập, hoặc đặt lại mật khẩu."),
    ("2", "Student", "Học viên cuối sử dụng hệ thống để học JLPT, luyện tập, thi thử, dùng AI OCR/Speech và theo dõi tiến độ."),
    ("3", "Staff", "Nhân viên nội dung quản lý học liệu, ngân hàng câu hỏi, bài quiz/exam, hỗ trợ học viên và chấm bài nói."),
    ("4", "StaffManager", "Quản lý nội dung, duyệt/từ chối/yêu cầu sửa, xuất bản/ẩn/lưu trữ nội dung do Staff gửi."),
    ("5", "Admin", "Quản trị viên hệ thống, quản lý tài khoản, cấu hình, báo cáo, notification rules và audit."),
    ("6", "External AI Service", "Dịch vụ OCR và Speech Recognition được gọi bất đồng bộ để trả kết quả gợi ý."),
    ("7", "SMTP Server", "Dịch vụ gửi email xác minh tài khoản, đặt lại mật khẩu và thông báo."),
]

actors_en = [
    ("1", "Guest", "A not-yet-logged-in user who can register a new account, log in, or reset a password."),
    ("2", "Student", "End user who studies JLPT content, practices, takes exams, uses OCR/Speech AI, and tracks progress."),
    ("3", "Staff", "Content operator who manages learning content, question banks, quizzes/exams, support tickets, and speaking grading."),
    ("4", "StaffManager", "Content manager who reviews, rejects, requests changes, publishes, hides, or archives Staff submissions."),
    ("5", "Admin", "System administrator who manages users, configuration, reports, notification rules, and audit logs."),
    ("6", "External AI Service", "Asynchronous OCR and Speech Recognition services returning suggested AI results."),
    ("7", "SMTP Server", "Email delivery service for verification, password reset, and notifications."),
]

use_cases_vi = [
    ("UC-01", "User Login", "Authentication", "Đăng nhập bằng email/password hoặc Google OAuth, nhận JWT và chuyển tới đúng dashboard theo role."),
    ("UC-02", "User Register", "Authentication", "Tạo tài khoản Student, gửi OTP xác minh email và kích hoạt tài khoản."),
    ("UC-03", "Reset Password", "Authentication", "Yêu cầu email đặt lại mật khẩu, xác minh token và cập nhật mật khẩu mới."),
    ("UC-04", "User Profile", "Student Account", "Xem và cập nhật hồ sơ cá nhân, avatar, số điện thoại, cấp độ JLPT mục tiêu."),
    ("UC-05", "Change Password", "Student Account", "Đổi mật khẩu khi đang đăng nhập sau khi xác thực mật khẩu hiện tại."),
    ("UC-06", "Learn Grammar", "Core Learning", "Học ngữ pháp theo JLPT level với cấu trúc, giải thích, ví dụ và tiến độ."),
    ("UC-07", "Learn Kanji", "Core Learning", "Học Kanji theo level, nghĩa, onyomi/kunyomi, số nét và ví dụ."),
    ("UC-08", "Learn Kana", "Core Learning", "Học Hiragana/Katakana với âm đọc và bài luyện nhận diện."),
    ("UC-09", "Vocabulary", "Core Learning", "Học từ vựng theo lesson/topic/level và thêm vào flashcard/notebook."),
    ("UC-10", "Take JLPT Mock Test", "Assessment", "Làm đề thi thử có thời gian; backend chấm điểm và lưu attempt mới."),
    ("UC-11", "Practice & Quiz", "Assessment", "Làm quiz theo chủ đề/bài học; backend tính điểm và trả kết quả."),
    ("UC-12", "Flashcard Learning", "SRS Review", "Ôn tập flashcard theo thuật toán spaced repetition và cập nhật lịch ôn."),
    ("UC-13", "Speaking Practice & AI Grading", "AI Skills", "Nộp audio luyện nói, nhận job_id, AI chấm gợi ý và Staff có thể override."),
    ("UC-14", "Reading Practice", "Skills Practice", "Luyện đọc hiểu theo cấp độ và lưu kết quả luyện tập."),
    ("UC-15", "Listening Practice", "Skills Practice", "Luyện nghe hiểu với audio và câu hỏi liên quan."),
    ("UC-16", "Dictionary & Search", "Learning Tools", "Tra cứu từ vựng/ngữ pháp/Kanji và lọc theo cấp độ."),
    ("UC-17", "Bookmark Learning", "Learning Tools", "Đánh dấu nội dung yêu thích hoặc cần ôn lại."),
    ("UC-18", "Logout", "Authentication", "Đăng xuất, thu hồi/huỷ token phiên hiện tại."),
    ("UC-19", "Learning Progress & Stats", "Analytics", "Xem tiến độ học, streak, tỷ lệ hoàn thành và kết quả luyện tập."),
    ("UC-20", "AI Handwriting Practice", "AI Skills", "Nộp ảnh/canvas Kanji viết tay, nhận similarity % bất đồng bộ."),
    ("UC-21", "View Student Progress", "Staff Student Management", "Staff xem tiến độ học viên phục vụ hỗ trợ và tư vấn."),
    ("UC-22", "Manage Student Accounts", "Staff Student Management", "Staff quản lý thông tin học viên trong phạm vi được phân quyền."),
    ("UC-23", "Suspend or Activate Account", "Staff Student Management", "Khoá/mở tài khoản học viên theo quy trình và audit."),
    ("UC-24", "Manage Question Bank", "Content Management", "Tạo/sửa/lưu nháp câu hỏi, đáp án, cấp độ, kỹ năng."),
    ("UC-25", "Manage Grammar Content", "Content Management", "Quản lý điểm ngữ pháp và gửi duyệt trước khi publish."),
    ("UC-26", "Manage Quiz", "Content Management", "Quản lý quiz; khi đã có attempt thì lock câu hỏi hoặc tạo version mới."),
    ("UC-27", "Manage Learning Content", "Content Management", "Quản lý course/lesson/Kanji/Kana/vocab/reading/listening/speaking."),
    ("UC-28", "Manage JLPT Mock Exams", "Content Management", "Tạo và quản lý đề thi thử JLPT theo cấp độ và kỹ năng."),
    ("UC-29", "Respond to Student Support", "Support", "Tiếp nhận và phản hồi ticket hỗ trợ của học viên."),
    ("UC-30", "Send Notifications", "Notification", "Gửi thông báo in-app/email tới học viên hoặc nhóm học viên."),
    ("UC-31", "Grade Speaking Submission", "Manual Grading", "Nghe audio, xem AI suggestion, nhập final_score và feedback."),
    ("UC-32", "View Quiz Results", "Analytics", "Xem kết quả quiz/mock test và phân tích câu hỏi."),
    ("UC-33", "Review Submitted Content", "Content Review", "StaffManager duyệt/từ chối/yêu cầu sửa nội dung đang pending_review."),
    ("UC-34", "Manage Published Content Status", "Content Review", "Publish/unpublish/archive/restore nội dung đã duyệt."),
    ("UC-35", "Login System", "Admin", "Admin đăng nhập vào khu vực quản trị bằng tài khoản được cấp."),
    ("UC-36", "View Dashboard", "Admin", "Xem tổng quan hệ thống, người dùng, nội dung, lượt học và bài thi."),
    ("UC-37", "User Management", "Admin", "Quản lý tài khoản Student/Staff/StaffManager/Admin và phân quyền."),
    ("UC-38", "Report Screen", "Admin", "Xem và xuất báo cáo học tập, sử dụng hệ thống và nội dung."),
    ("UC-39", "Settings", "Admin", "Cấu hình hệ thống, SMTP, session, maintenance và tham số vận hành."),
    ("UC-40", "Notification Rules", "Admin", "Quản lý rule/template thông báo tự động theo sự kiện."),
]

use_cases_en = [
    ("UC-01", "User Login", "Authentication", "Authenticate by email/password or Google OAuth, issue JWT, and route to the role dashboard."),
    ("UC-02", "User Register", "Authentication", "Create a Student account, send email OTP, and activate the account."),
    ("UC-03", "Reset Password", "Authentication", "Send a reset email, validate the reset token, and update the password."),
    ("UC-04", "User Profile", "Student Account", "View and update personal profile, avatar, phone number, and target JLPT level."),
    ("UC-05", "Change Password", "Student Account", "Change password after validating the current password."),
    ("UC-06", "Learn Grammar", "Core Learning", "Study JLPT grammar with patterns, explanation, examples, and progress tracking."),
    ("UC-07", "Learn Kanji", "Core Learning", "Study Kanji by level, meaning, onyomi/kunyomi, stroke count, and examples."),
    ("UC-08", "Learn Kana", "Core Learning", "Study Hiragana/Katakana with pronunciation and recognition practice."),
    ("UC-09", "Vocabulary", "Core Learning", "Study vocabulary by lesson/topic/level and add items to flashcards/notebook."),
    ("UC-10", "Take JLPT Mock Test", "Assessment", "Take timed mock exams; backend grades and stores a new attempt."),
    ("UC-11", "Practice & Quiz", "Assessment", "Take lesson/topic quizzes; backend calculates score and returns results."),
    ("UC-12", "Flashcard Learning", "SRS Review", "Review flashcards using spaced repetition and update the next review schedule."),
    ("UC-13", "Speaking Practice & AI Grading", "AI Skills", "Submit speaking audio, receive a job_id, get AI suggestions, and allow Staff override."),
    ("UC-14", "Reading Practice", "Skills Practice", "Practice reading comprehension by level and store results."),
    ("UC-15", "Listening Practice", "Skills Practice", "Practice listening comprehension with audio and linked questions."),
    ("UC-16", "Dictionary & Search", "Learning Tools", "Search vocabulary, grammar, and Kanji with level filters."),
    ("UC-17", "Bookmark Learning", "Learning Tools", "Bookmark favorite or review-needed learning content."),
    ("UC-18", "Logout", "Authentication", "Log out and revoke/invalidate the current session token."),
    ("UC-19", "Learning Progress & Stats", "Analytics", "View progress, streaks, completion rate, and practice results."),
    ("UC-20", "AI Handwriting Practice", "AI Skills", "Submit handwritten Kanji image/canvas and receive asynchronous similarity %."),
    ("UC-21", "View Student Progress", "Staff Student Management", "View student progress for support and advising."),
    ("UC-22", "Manage Student Accounts", "Staff Student Management", "Manage student information within delegated permissions."),
    ("UC-23", "Suspend or Activate Account", "Staff Student Management", "Suspend or reactivate student accounts with audit trail."),
    ("UC-24", "Manage Question Bank", "Content Management", "Create/edit draft questions, answers, levels, and skills."),
    ("UC-25", "Manage Grammar Content", "Content Management", "Manage grammar content and submit it for review before publishing."),
    ("UC-26", "Manage Quiz", "Content Management", "Manage quizzes; lock attempted questions or create a new version."),
    ("UC-27", "Manage Learning Content", "Content Management", "Manage courses, lessons, Kanji, Kana, vocabulary, reading, listening, and speaking content."),
    ("UC-28", "Manage JLPT Mock Exams", "Content Management", "Create and manage JLPT mock exams by level and skill."),
    ("UC-29", "Respond to Student Support", "Support", "Receive and respond to student support tickets."),
    ("UC-30", "Send Notifications", "Notification", "Send in-app/email notifications to students or student groups."),
    ("UC-31", "Grade Speaking Submission", "Manual Grading", "Review audio, AI suggestions, final_score, and feedback."),
    ("UC-32", "View Quiz Results", "Analytics", "View quiz/mock exam results and question analytics."),
    ("UC-33", "Review Submitted Content", "Content Review", "Approve, reject, or request changes for pending_review content."),
    ("UC-34", "Manage Published Content Status", "Content Review", "Publish, unpublish, archive, or restore approved content."),
    ("UC-35", "Login System", "Admin", "Admin signs in to the administration area with a provisioned account."),
    ("UC-36", "View Dashboard", "Admin", "View system overview, users, content, activity, and assessments."),
    ("UC-37", "User Management", "Admin", "Manage Student, Staff, StaffManager, and Admin accounts and permissions."),
    ("UC-38", "Report Screen", "Admin", "View and export learning, system usage, and content reports."),
    ("UC-39", "Settings", "Admin", "Configure system settings, SMTP, sessions, maintenance, and operations parameters."),
    ("UC-40", "Notification Rules", "Admin", "Manage automated notification rules/templates by event."),
]

business_rules_vi = [
    ("BR-01", "Authorization phải kiểm tra cả Role và Subscription/Level; UI chỉ là UX, backend vẫn trả 401/403."),
    ("BR-02", "Điểm quiz/mock exam chỉ được tính tại backend Service layer; client không gửi score."),
    ("BR-03", "Score luôn nằm trong khoảng 0..max_score và mỗi lần nộp tạo attempt mới."),
    ("BR-04", "Attempt đã SUBMITTED là bất biến; không sửa điểm hoặc đáp án sau khi nộp."),
    ("BR-05", "Quiz/Exam đã có attempt thì câu hỏi bị lock; sửa nội dung phải tạo version mới."),
    ("BR-06", "Bài học tiếp theo chỉ mở khi bài trước hoàn thành; user_progress chỉ tăng."),
    ("BR-07", "Nội dung Staff tạo phải qua StaffManager review trước khi publish."),
    ("BR-08", "Nội dung VIP chỉ hiển thị khi subscription VIP còn hiệu lực; cache subscription tối đa 5 phút."),
    ("BR-09", "AI task chạy async, trả job_id, timeout + retry tối đa 3 lần + fallback rõ ràng."),
    ("BR-10", "File ảnh/audio lưu ở /uploads hoặc S3; không lưu BLOB trong DB."),
    ("BR-11", "Soft delete bắt buộc cho dữ liệu quan trọng; không hard delete."),
    ("BR-12", "Thao tác quan trọng của Admin/Staff/StaffManager phải ghi audit log."),
]

business_rules_en = [
    ("BR-01", "Authorization must check both Role and Subscription/Level; UI hiding is only UX and backend still returns 401/403."),
    ("BR-02", "Quiz/mock exam scores are calculated only in the backend Service layer; clients never submit score."),
    ("BR-03", "Score must be within 0..max_score and every submission creates a new attempt."),
    ("BR-04", "SUBMITTED attempts are immutable; score and answers cannot be modified after submission."),
    ("BR-05", "Once a quiz/exam has attempts, its questions are locked; changes require a new version."),
    ("BR-06", "The next lesson unlocks only after the previous lesson is completed; user_progress only increases."),
    ("BR-07", "Staff-created content must be reviewed by StaffManager before publishing."),
    ("BR-08", "VIP content is visible only when the VIP subscription is active; subscription cache is at most 5 minutes."),
    ("BR-09", "AI tasks run asynchronously, return job_id, and use timeout + up to 3 retries + clear fallback."),
    ("BR-10", "Image/audio files are stored in /uploads or S3; BLOB storage in DB is not allowed."),
    ("BR-11", "Soft delete is mandatory for important data; hard delete is forbidden."),
    ("BR-12", "Important Admin/Staff/StaffManager actions must be recorded in audit logs."),
]

entities_vi = [
    ("1", "student_users", "Tài khoản học viên, OAuth, cấp độ JLPT, subscription và thống kê học tập."),
    ("2", "staff_users", "Tài khoản Staff và StaffManager, vai trò nội bộ, trạng thái."),
    ("3", "admin_users", "Tài khoản quản trị hệ thống."),
    ("4", "auth_tokens", "Session, refresh token, email verification, password reset."),
    ("5", "courses", "Khóa học JLPT theo level, trạng thái, VIP flag."),
    ("6", "lessons", "Bài học thuộc course, gồm lesson/reading/listening/speaking."),
    ("7", "kanji / kana_characters / vocabulary / grammar_points", "Nội dung học cốt lõi."),
    ("8", "questions / assessments / question_assignments", "Ngân hàng câu hỏi, quiz và đề thi thử."),
    ("9", "test_attempts / attempt_answers", "Lịch sử làm bài và câu trả lời."),
    ("10", "student_submissions", "Bài speaking/handwriting, AI score suggestion và final score."),
    ("11", "student_content_progress / flashcards", "Tiến độ học, bookmark và SRS flashcard."),
    ("12", "tickets / ticket_replies / notifications / system_settings / admin_audit_logs", "Hỗ trợ, thông báo, cấu hình và audit."),
]

entities_en = [
    ("1", "student_users", "Student accounts, OAuth, JLPT level, subscription, and learning statistics."),
    ("2", "staff_users", "Staff and StaffManager accounts, internal role, and status."),
    ("3", "admin_users", "System administrator accounts."),
    ("4", "auth_tokens", "Session, refresh token, email verification, and password reset."),
    ("5", "courses", "JLPT courses by level, status, and VIP flag."),
    ("6", "lessons", "Course lessons including lesson/reading/listening/speaking."),
    ("7", "kanji / kana_characters / vocabulary / grammar_points", "Core learning content."),
    ("8", "questions / assessments / question_assignments", "Question bank, quizzes, and mock exams."),
    ("9", "test_attempts / attempt_answers", "Assessment attempt history and answers."),
    ("10", "student_submissions", "Speaking/handwriting submissions, AI score suggestion, and final score."),
    ("11", "student_content_progress / flashcards", "Learning progress, bookmarks, and SRS flashcards."),
    ("12", "tickets / ticket_replies / notifications / system_settings / admin_audit_logs", "Support, notifications, settings, and audit."),
]


def context_diagram(lang):
    if lang == "vi":
        return """```mermaid
flowchart LR
  Student[Student] --> FE[React Web App]
  Staff[Staff] --> FE
  Manager[StaffManager] --> FE
  Admin[Admin] --> FE
  FE --> API[Spring Boot REST API]
  API --> DB[(MySQL 8)]
  API --> FS[/Uploads or S3/]
  API --> AI[OCR/Speech AI Service]
  API --> SMTP[SMTP Email Server]
```"""
    return """```mermaid
flowchart LR
  Student[Student] --> FE[React Web App]
  Staff[Staff] --> FE
  Manager[StaffManager] --> FE
  Admin[Admin] --> FE
  FE --> API[Spring Boot REST API]
  API --> DB[(MySQL 8)]
  API --> FS[/Uploads or S3/]
  API --> AI[OCR/Speech AI Service]
  API --> SMTP[SMTP Email Server]
```"""


def screens_flow(lang):
    labels = {
        "vi": ("Trang chủ", "Đăng nhập", "Dashboard Student", "Học nội dung", "Quiz/Mock Test", "AI Practice", "Staff", "Manager", "Admin"),
        "en": ("Home", "Login", "Student Dashboard", "Learning Content", "Quiz/Mock Test", "AI Practice", "Staff", "Manager", "Admin"),
    }[lang]
    return f"""```mermaid
flowchart TD
  A[{labels[0]}] --> B[{labels[1]}]
  B --> C[{labels[2]}]
  C --> D[{labels[3]}]
  C --> E[{labels[4]}]
  C --> F[{labels[5]}]
  B --> G[{labels[6]}]
  B --> H[{labels[7]}]
  B --> I[{labels[8]}]
```"""


def uc_diagram(lang):
    if lang == "vi":
        return """```mermaid
flowchart LR
  Student((Student)) --> S1[UC-01..20 Học, luyện tập, thi, AI]
  Staff((Staff)) --> T1[UC-21..32 Quản lý nội dung, hỗ trợ, chấm bài]
  Manager((StaffManager)) --> M1[UC-33..34 Duyệt và xuất bản nội dung]
  Admin((Admin)) --> A1[UC-35..40 Quản trị hệ thống]
```"""
    return """```mermaid
flowchart LR
  Student((Student)) --> S1[UC-01..20 Learning, practice, exams, AI]
  Staff((Staff)) --> T1[UC-21..32 Content, support, grading]
  Manager((StaffManager)) --> M1[UC-33..34 Review and publishing]
  Admin((Admin)) --> A1[UC-35..40 System administration]
```"""


screen_auth = [
    ("Public", "X", "X", "X", "X", "Login/register/forgot/reset/home; public API only for /api/auth/*"),
    ("Student dashboard & learning", "X", "", "", "", "Requires STUDENT role plus subscription/level checks."),
    ("Quiz, mock test, flashcard, AI practice", "X", "", "", "", "Score/time/AI validation handled server-side."),
    ("Staff dashboard/content/questions/assessments/grading", "", "X", "", "", "Requires STAFF role."),
    ("Manager review queue/content pipeline", "", "", "X", "X", "Requires StaffManager or Admin; service layer validates staff_role."),
    ("Admin users/settings/reports/notification rules", "", "", "", "X", "Requires ADMIN role."),
]

non_ui_vi = [
    ("1", "Authentication", "AuthEventListener", "Gửi email xác minh tài khoản và reset password bất đồng bộ."),
    ("2", "Speaking", "SpeakingAsyncProcessor", "Xử lý bài nộp speaking sau khi nhận request; trả job/status để frontend poll."),
    ("3", "Notification", "NotificationDispatcher", "Fan-out thông báo in-app/email và retry email trong outbox."),
    ("4", "Assessment", "Quiz/Exam scoring service", "Tính điểm, validate thời gian, tạo attempt mới, ghi audit."),
    ("5", "Subscription", "Subscription validation service", "Kiểm tra quyền VIP real-time/cache tối đa 5 phút."),
]

non_ui_en = [
    ("1", "Authentication", "AuthEventListener", "Sends verification and password reset emails asynchronously."),
    ("2", "Speaking", "SpeakingAsyncProcessor", "Processes speaking submissions after request acceptance; exposes job/status polling."),
    ("3", "Notification", "NotificationDispatcher", "Fans out in-app/email notifications and retries email outbox delivery."),
    ("4", "Assessment", "Quiz/Exam scoring service", "Calculates score, validates time, creates new attempts, and writes audit."),
    ("5", "Subscription", "Subscription validation service", "Validates VIP access in real time with at most 5-minute cache."),
]

messages_vi = [
    ("1", "MSG01", "Inline", "No search result", "Không có kết quả phù hợp."),
    ("2", "MSG02", "Field error", "Required field is empty", "Trường này là bắt buộc."),
    ("3", "MSG03", "Toast", "Save success", "Lưu dữ liệu thành công."),
    ("4", "MSG04", "Toast", "Delete/soft delete success", "Cập nhật trạng thái thành công."),
    ("5", "MSG05", "Inline", "Invalid login", "Email hoặc mật khẩu không đúng."),
    ("6", "MSG06", "Inline", "Forbidden", "Bạn không có quyền truy cập chức năng này."),
    ("7", "MSG07", "Toast", "AI job accepted", "Bài nộp đã được tiếp nhận và đang xử lý."),
    ("8", "MSG08", "Toast", "AI failed", "Không thể xử lý AI lúc này. Vui lòng thử lại sau."),
]

messages_en = [
    ("1", "MSG01", "Inline", "No search result", "No matching results."),
    ("2", "MSG02", "Field error", "Required field is empty", "This field is required."),
    ("3", "MSG03", "Toast", "Save success", "Data saved successfully."),
    ("4", "MSG04", "Toast", "Delete/soft delete success", "Status updated successfully."),
    ("5", "MSG05", "Inline", "Invalid login", "Email or password is incorrect."),
    ("6", "MSG06", "Inline", "Forbidden", "You do not have permission to access this feature."),
    ("7", "MSG07", "Toast", "AI job accepted", "Your submission has been accepted and is being processed."),
    ("8", "MSG08", "Toast", "AI failed", "AI processing is unavailable now. Please try again later."),
]


def table_md(headers, rows):
    s = "| " + " | ".join(headers) + " |\n"
    s += "| " + " | ".join(["---"] * len(headers)) + " |\n"
    for row in rows:
        s += "| " + " | ".join(str(x).replace("\n", "<br>") for x in row) + " |\n"
    return s


def make_markdown(lang):
    vi = lang == "vi"
    title = "ĐẶC TẢ YÊU CẦU PHẦN MỀM" if vi else "SOFTWARE REQUIREMENT SPECIFICATION"
    project = "Japanese Skill Practice Platform"
    actors = actors_vi if vi else actors_en
    ucs = use_cases_vi if vi else use_cases_en
    brs = business_rules_vi if vi else business_rules_en
    entities = entities_vi if vi else entities_en
    non_ui = non_ui_vi if vi else non_ui_en
    messages = messages_vi if vi else messages_en

    if vi:
        intro = "Japanese Skill Practice Platform là hệ thống web học và luyện thi tiếng Nhật JLPT từ N5 đến N1, gồm React 18 frontend, Spring Boot 3/Java 21 backend và MySQL 8. Hệ thống cung cấp học liệu Kanji, Kana, từ vựng, ngữ pháp, quiz, mock exam, flashcard SRS, luyện nói, luyện viết tay với AI OCR/Speech, quản trị nội dung và báo cáo."
        main_process = [
            ("1.2.1", "Xác thực và vào hệ thống", "Người dùng đăng ký/đăng nhập, hệ thống xác thực JWT và chuyển tới dashboard theo role."),
            ("1.2.2", "Học tập và luyện thi", "Student học nội dung theo JLPT level, làm quiz/mock exam, backend tính điểm và cập nhật tiến độ."),
            ("1.2.3", "AI Skills", "Student nộp ảnh/audio, backend tạo AI job bất đồng bộ, gọi OCR/Speech, lưu gợi ý điểm và trả kết quả."),
            ("1.2.4", "Quản lý và duyệt nội dung", "Staff tạo nội dung draft, StaffManager duyệt/publish, hệ thống audit mọi thao tác quan trọng."),
            ("1.2.5", "Quản trị hệ thống", "Admin quản lý tài khoản, cấu hình, notification rules, subscription và báo cáo."),
        ]
    else:
        intro = "Japanese Skill Practice Platform is a web-based JLPT learning and exam practice system for N5 to N1, built with React 18, Spring Boot 3/Java 21, and MySQL 8. It provides Kanji, Kana, vocabulary, grammar, quizzes, mock exams, SRS flashcards, speaking and handwriting practice with OCR/Speech AI, content governance, and reporting."
        main_process = [
            ("1.2.1", "Authentication and Entry", "Users register/sign in, the system validates JWT, and routes them to the correct role dashboard."),
            ("1.2.2", "Learning and Assessment", "Students study JLPT-level content, take quizzes/mock exams, and backend calculates score and progress."),
            ("1.2.3", "AI Skills", "Students submit image/audio, backend creates asynchronous AI jobs, calls OCR/Speech, stores suggestions, and returns results."),
            ("1.2.4", "Content Management and Review", "Staff create draft content, StaffManager reviews/publishes it, and the system audits important actions."),
            ("1.2.5", "System Administration", "Admin manages accounts, configuration, notification rules, subscription, and reports."),
        ]

    md = f"""# {title}

**Project Name:** {project}  
**Version:** 1.0  
**Location/Date:** Hanoi, July 2026  
**Source Template:** `Guides  Templates-20260721/Template1_SRS Document.md` and `Template1_SRS Document.docx`

## Table of Contents

# I. Record of Changes

| Date | A/M/D | In charge | Change Description |
| --- | --- | --- | --- |
| 2026-07-23 | A | AI Agent | Initial SRS document generated in {'Vietnamese' if vi else 'English'} according to Template1 structure. |

*A - Added, M - Modified, D - Deleted*

# II. Software Requirement Specification

## 1. Overall Requirements

### 1.1 Context Diagram

{intro}

{context_diagram(lang)}

### 1.2 Main Business Processes

{table_md(['#', 'Business Process', 'Description'], main_process)}

### 1.3 User Requirements

#### 1.3.1 Actors

{table_md(['#', 'Actor', 'Description'], actors)}

#### 1.3.2 Use Cases (UC)

{table_md(['ID', 'Use Case', 'Feature', 'Use Case Description'], ucs)}

#### 1.3.3 Use Case Diagrams

{uc_diagram(lang)}

### 1.4 System Functionalities

#### 1.4.1 Screens Flow

{screens_flow(lang)}

#### 1.4.2 Screen Authorization

{table_md(['Screen', 'Student', 'Staff', 'StaffManager', 'Admin', 'Notes'], screen_auth)}

#### 1.4.3 Non-UI Functions

{table_md(['#', 'Feature', 'System Function', 'Description'], non_ui)}

### 1.5 Entity Relationship Diagram

```mermaid
erDiagram
  STUDENT_USERS ||--o{{ TEST_ATTEMPTS : submits
  STUDENT_USERS ||--o{{ STUDENT_CONTENT_PROGRESS : tracks
  STUDENT_USERS ||--o{{ STUDENT_SUBMISSIONS : uploads
  COURSES ||--o{{ LESSONS : contains
  LESSONS ||--o{{ VOCABULARY : includes
  LESSONS ||--o{{ GRAMMAR_POINTS : includes
  ASSESSMENTS ||--o{{ QUESTION_ASSIGNMENTS : contains
  QUESTIONS ||--o{{ QUESTION_ASSIGNMENTS : assigned
  TEST_ATTEMPTS ||--o{{ ATTEMPT_ANSWERS : records
  STAFF_USERS ||--o{{ ADMIN_AUDIT_LOGS : performs
```

**Entities Description**

{table_md(['#', 'Entity', 'Description'], entities)}

## 2. Use Case Specifications

### 2.1 Authentication

#### 2.1.1 UC-01 User Login

| Primary Actors | Student, Staff, StaffManager, Admin | Secondary Actors | Google OAuth Provider |
| --- | --- | --- | --- |
| Description | {'Người dùng đăng nhập để truy cập các chức năng được cá nhân hóa và được phân quyền.' if vi else 'The user signs in to access personalized and authorized system features.'} |  |  |
| Preconditions | {'Tài khoản tồn tại, hoạt động và không bị xóa mềm.' if vi else 'The account exists, is active, and is not soft-deleted.'} |  |  |
| Postconditions | {'JWT/refresh token hợp lệ được cấp, login event được ghi nhận, người dùng vào đúng dashboard.' if vi else 'Valid JWT/refresh token is issued, login event is tracked, and the user reaches the correct dashboard.'} |  |  |
| Normal Sequence/Flow | {'1. Người dùng mở màn hình đăng nhập. 2. Nhập email và mật khẩu. 3. Hệ thống validate input và xác thực tài khoản. 4. Hệ thống cấp token. 5. Frontend chuyển hướng theo role.' if vi else '1. User opens login screen. 2. User enters email and password. 3. System validates input and authenticates the account. 4. System issues tokens. 5. Frontend routes by role.'} |  |  |
| Alternative Sequences/Flows | {'Google OAuth; sai mật khẩu; tài khoản chưa xác minh; tài khoản bị khóa; thiếu quyền truy cập route.' if vi else 'Google OAuth; wrong password; unverified account; locked account; route access denied.'} |  |  |

#### 2.1.2 UC-02 User Register

| Primary Actors | Student | Secondary Actors | SMTP Server |
| --- | --- | --- | --- |
| Description | {'Khách tạo tài khoản Student và xác minh email bằng OTP.' if vi else 'A guest creates a Student account and verifies email with OTP.'} |  |  |
| Preconditions | {'Email chưa tồn tại trong hệ thống.' if vi else 'Email is not already registered.'} |  |  |
| Postconditions | {'Tài khoản Student được kích hoạt sau xác minh email.' if vi else 'Student account is activated after email verification.'} |  |  |
| Normal Sequence/Flow | {'1. Nhập thông tin đăng ký. 2. Backend validate và tạo user. 3. Tạo OTP 6 chữ số. 4. Gửi email. 5. Người dùng xác minh OTP.' if vi else '1. Enter registration data. 2. Backend validates and creates user. 3. Generate 6-digit OTP. 4. Send email. 5. User verifies OTP.'} |  |  |
| Alternative Sequences/Flows | {'Email trùng, OTP sai/hết hạn, gửi lại OTP bị rate limit.' if vi else 'Duplicate email, wrong/expired OTP, resend rate limit.'} |  |  |

### 2.2 Assessment

#### 2.2.1 UC-10 Take JLPT Mock Test

| Primary Actors | Student | Secondary Actors | None |
| --- | --- | --- | --- |
| Description | {'Student làm đề thi thử JLPT có thời gian và nhận kết quả do backend chấm.' if vi else 'Student takes a timed JLPT mock exam and receives backend-graded results.'} |  |  |
| Preconditions | {'Student đã đăng nhập, có quyền truy cập level/subscription, đề thi ở trạng thái published.' if vi else 'Student is signed in, has level/subscription access, and the exam is published.'} |  |  |
| Postconditions | {'Một bản ghi attempt mới được tạo với score hợp lệ và answers bất biến.' if vi else 'A new attempt is created with valid score and immutable answers.'} |  |  |
| Normal Sequence/Flow | {'1. Chọn đề. 2. Backend tạo/ghi nhận phiên làm bài. 3. Student trả lời. 4. Nộp bài. 5. Backend validate thời gian, tính điểm, lưu attempt, trả kết quả.' if vi else '1. Select exam. 2. Backend creates/tracks attempt session. 3. Student answers. 4. Submit. 5. Backend validates time, calculates score, stores attempt, returns result.'} |  |  |
| Alternative Sequences/Flows | {'Hết thời gian, thiếu quyền VIP, đề đã bị unpublish, request submit lặp.' if vi else 'Time expired, VIP access missing, exam unpublished, duplicate submit request.'} |  |  |

### 2.3 AI Skills

#### 2.3.1 UC-13 Speaking Practice & AI Grading

| Primary Actors | Student | Secondary Actors | AI Speech Service, Staff |
| --- | --- | --- | --- |
| Description | {'Student nộp audio luyện nói, AI đưa điểm gợi ý, Staff có thể chấm cuối.' if vi else 'Student submits speaking audio, AI returns suggested score, and Staff can finalize grading.'} |  |  |
| Preconditions | {'Student đã đăng nhập, bài speaking published, file audio hợp lệ.' if vi else 'Student is signed in, speaking lesson is published, and audio file is valid.'} |  |  |
| Postconditions | {'Submission được lưu, trạng thái AI được cập nhật, final_score chỉ có sau Staff/logic được phép.' if vi else 'Submission is stored, AI status is updated, and final_score is set only by authorized Staff/logic.'} |  |  |
| Normal Sequence/Flow | {'1. Upload audio. 2. Backend lưu file và tạo job. 3. Trả job_id. 4. Processor gọi AI với timeout/retry. 5. Student poll kết quả. 6. Staff review nếu cần.' if vi else '1. Upload audio. 2. Backend stores file and creates job. 3. Return job_id. 4. Processor calls AI with timeout/retry. 5. Student polls result. 6. Staff reviews when needed.'} |  |  |
| Alternative Sequences/Flows | {'AI timeout/failure, file không hợp lệ, Staff override điểm AI.' if vi else 'AI timeout/failure, invalid file, Staff overrides AI score.'} |  |  |

### 2.4 Content Review

#### 2.4.1 UC-33 Review Submitted Content

| Primary Actors | StaffManager | Secondary Actors | Staff |
| --- | --- | --- | --- |
| Description | {'StaffManager duyệt nội dung Staff gửi trước khi học viên thấy.' if vi else 'StaffManager reviews Staff-submitted content before students can see it.'} |  |  |
| Preconditions | {'Nội dung ở trạng thái pending_review và StaffManager có quyền hợp lệ.' if vi else 'Content is pending_review and StaffManager has valid permission.'} |  |  |
| Postconditions | {'Nội dung được approve/reject/request changes, có audit log và thông báo tới Staff.' if vi else 'Content is approved/rejected/requested for changes, with audit log and Staff notification.'} |  |  |
| Normal Sequence/Flow | {'1. Mở review queue. 2. Lọc và mở chi tiết. 3. Kiểm tra snapshot/nội dung. 4. Chọn Approve/Reject/Request Changes. 5. Hệ thống cập nhật trạng thái và audit.' if vi else '1. Open review queue. 2. Filter/open details. 3. Inspect snapshot/content. 4. Choose Approve/Reject/Request Changes. 5. System updates status and audit.'} |  |  |
| Alternative Sequences/Flows | {'Nội dung đã được người khác xử lý, thiếu dữ liệu bắt buộc, không đủ quyền.' if vi else 'Content already handled by another reviewer, required data missing, insufficient permission.'} |  |  |

## 3. Functional Requirements

### 3.1 Authentication

#### 3.1.1 Login Screen

{'Cho phép người dùng nhập email/password, đăng nhập Google, đi tới forgot password/register. Backend trả token và role; frontend chỉ điều hướng theo response.' if vi else 'Allows users to enter email/password, use Google login, and navigate to forgot password/register. Backend returns token and role; frontend only routes from response.'}

| Field Name | Description |
| --- | --- |
| Email | {'Bắt buộc, định dạng email, tối đa 255 ký tự.' if vi else 'Required, email format, maximum 255 characters.'} |
| Password | {'Bắt buộc, không log plaintext.' if vi else 'Required, plaintext is never logged.'} |
| Login button | {'Gọi API đăng nhập và hiển thị loading/error.' if vi else 'Calls login API and shows loading/error.'} |
| Google login | {'Khởi tạo OAuth flow nếu cấu hình sẵn.' if vi else 'Starts OAuth flow when configured.'} |

### 3.2 Student Learning

#### 3.2.1 Course/Lesson Detail

{'Hiển thị bài học theo JLPT level, lesson order, trạng thái hoàn thành, VIP lock và nội dung Kanji/Kana/Vocab/Grammar liên quan.' if vi else 'Displays lessons by JLPT level, lesson order, completion state, VIP lock, and related Kanji/Kana/Vocab/Grammar content.'}

| Field Name | Description |
| --- | --- |
| JLPT level | N5, N4, N3, N2, N1 |
| Lesson status | LOCKED, AVAILABLE, COMPLETED |
| Content blocks | Text, examples, image/audio path, linked questions |
| Complete action | {'Gửi completion event, backend cập nhật progress.' if vi else 'Sends completion event; backend updates progress.'} |

### 3.3 Assessment

#### 3.3.1 Quiz/Mock Test Attempt

{'Màn hình làm bài chỉ gửi answer payload; mọi logic điểm, thời gian và attempt nằm ở backend.' if vi else 'Attempt screen sends only answer payload; scoring, timing, and attempt records belong to backend.'}

| Field Name | Description |
| --- | --- |
| Question list | {'Câu hỏi theo assessment/question assignment đã publish.' if vi else 'Questions from published assessment/question assignment.'} |
| Timer | {'Hiển thị thời gian UX; backend validate thời gian thật.' if vi else 'UX timer; backend validates actual time.'} |
| Submit button | {'Gửi answers, không gửi score.' if vi else 'Submits answers, never score.'} |
| Result view | {'Hiển thị score, max_score, đúng/sai sau khi backend trả kết quả.' if vi else 'Shows score, max_score, correct/wrong after backend returns result.'} |

### 3.4 Staff / Manager / Admin

#### 3.4.1 Content Management and Review

{'Staff tạo nội dung draft/pending_review; StaffManager duyệt và publish; Admin có quyền cấu hình/quản trị. Các thao tác quan trọng đều audit.' if vi else 'Staff creates draft/pending_review content; StaffManager reviews and publishes; Admin configures/administers. Important actions are audited.'}

| Field Name | Description |
| --- | --- |
| Status | draft, pending_review, published, archived |
| Review action | Approve, Reject, Request Changes |
| Audit data | actor, action, target, reason, timestamp |

## 4. Non-Functional Requirements

### 4.1 External Interfaces

- REST API prefix: `/api/[resource]`.
- JSON response format: `{{ "status": number, "message": string, "data": object }}`.
- Database: MySQL 8, UTF-8/utf8mb4, UTC timestamps.
- File storage: `/uploads` in development or S3-compatible storage in production.
- AI OCR/Speech and SMTP are external services with fallback/retry behavior.

### 4.2 Quality Attributes

#### 4.2.1 Usability

{'Giao diện web tiếng Việt, responsive, có trạng thái loading/error/empty rõ ràng và không dùng frontend để thay thế validation nghiệp vụ.' if vi else 'Vietnamese-first responsive web UI with clear loading/error/empty states; frontend validation never replaces business validation.'}

#### 4.2.2 Performance

{'API phổ biến phản hồi trung bình dưới 2 giây trong điều kiện bình thường; AI trả job_id ngay và xử lý nền; email dùng queue/retry.' if vi else 'Common APIs respond within 2 seconds on average under normal conditions; AI returns job_id immediately and runs in background; email uses queue/retry.'}

#### 4.2.3 Security

{'JWT bắt buộc cho API private, bcrypt cost tối thiểu 10, DTO pattern bắt buộc, không expose Entity, không hardcode secret.' if vi else 'JWT is required for private APIs, bcrypt cost is at least 10, DTO pattern is mandatory, Entities are not exposed, and secrets are not hardcoded.'}

#### 4.2.4 Reliability and Auditability

{'Soft delete, audit log, immutable submitted attempts, retry/fallback cho AI/email, global exception handler trả JSON chuẩn.' if vi else 'Soft delete, audit logs, immutable submitted attempts, retry/fallback for AI/email, and global exception handler with standard JSON responses.'}

## 5. Requirement Appendix

### 5.1 Business Rules

{table_md(['ID', 'Rule Definition'], brs)}

### 5.2 System Messages

{table_md(['#', 'Message code', 'Message Type', 'Context', 'Content'], messages)}

### 5.3 Other Requirements

- {'Không lưu secrets/password/API keys trong source control.' if vi else 'Do not store secrets/passwords/API keys in source control.'}
- {'Không thay đổi schema DB nếu không có migration.' if vi else 'Do not change database schema without migration.'}
- {'Không đặt business logic tính điểm, phân quyền hoặc subscription ở frontend.' if vi else 'Do not place scoring, authorization, or subscription business logic in frontend.'}
- {'Tài liệu nguồn tham chiếu: AGENTS.md, CLAUDE.md, docs/01-SRS-Requirements/shared_context.md, docs/01-SRS-Requirements/use-cases/Bao_cao_dac_ta_Use_Case.md.' if vi else 'Reference sources: AGENTS.md, CLAUDE.md, docs/01-SRS-Requirements/shared_context.md, docs/01-SRS-Requirements/use-cases/Bao_cao_dac_ta_Use_Case.md.'}
"""
    return md


def clear_document(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(str(h))
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def add_para(doc, text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def make_docx(lang, path):
    vi = lang == "vi"
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title = "ĐẶC TẢ YÊU CẦU PHẦN MỀM" if vi else "SOFTWARE REQUIREMENT SPECIFICATION"
    actors = actors_vi if vi else actors_en
    ucs = use_cases_vi if vi else use_cases_en
    brs = business_rules_vi if vi else business_rules_en
    entities = entities_vi if vi else entities_en
    non_ui = non_ui_vi if vi else non_ui_en
    messages = messages_vi if vi else messages_en

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Japanese Skill Practice Platform")
    r.bold = True
    r.font.size = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Hanoi, July 2026")
    doc.add_paragraph()

    add_heading(doc, "Table of Contents", 1)
    for line in [
        "I. Record of Changes",
        "II. Software Requirement Specification",
        "1. Overall Requirements",
        "2. Use Case Specifications",
        "3. Functional Requirements",
        "4. Non-Functional Requirements",
        "5. Requirement Appendix",
    ]:
        add_para(doc, line)

    add_heading(doc, "I. Record of Changes", 1)
    add_table(doc, ["Date", "A/M/D", "In charge", "Change Description"], [("2026-07-23", "A", "AI Agent", f"Initial SRS document generated in {'Vietnamese' if vi else 'English'} according to Template1.")])
    add_para(doc, "*A - Added, M - Modified, D - Deleted")

    add_heading(doc, "II. Software Requirement Specification", 1)
    add_heading(doc, "1. Overall Requirements", 2)
    add_heading(doc, "1.1 Context Diagram", 3)
    add_para(doc, "Japanese Skill Practice Platform " + ("là hệ thống web học và luyện thi tiếng Nhật JLPT từ N5 đến N1, tích hợp học liệu, assessment, SRS flashcard, AI OCR/Speech, quản trị nội dung và báo cáo." if vi else "is a web-based JLPT learning and exam practice system from N5 to N1, integrating learning content, assessment, SRS flashcards, OCR/Speech AI, content governance, and reporting."))
    add_para(doc, "Context: Student/Staff/StaffManager/Admin use React Web App; the app calls Spring Boot REST API; backend integrates MySQL 8, file storage, OCR/Speech AI service, and SMTP email server.")

    add_heading(doc, "1.2 Main Business Processes", 3)
    processes = [
        ("1", "Authentication and role routing", "Đăng ký/đăng nhập, JWT, route theo role." if vi else "Register/sign in, issue JWT, and route by role."),
        ("2", "Learning and assessment", "Học theo JLPT level, làm quiz/mock exam, backend tính điểm." if vi else "Study by JLPT level, take quiz/mock exam, backend calculates score."),
        ("3", "AI skills", "Nộp ảnh/audio, tạo job bất đồng bộ, poll kết quả." if vi else "Submit image/audio, create asynchronous job, poll result."),
        ("4", "Content review", "Staff tạo, StaffManager duyệt/publish, audit." if vi else "Staff creates, StaffManager reviews/publishes, audit is recorded."),
        ("5", "Administration", "Admin quản lý user, settings, subscription, report." if vi else "Admin manages users, settings, subscription, and reports."),
    ]
    add_table(doc, ["#", "Business Process", "Description"], processes)

    add_heading(doc, "1.3 User Requirements", 3)
    add_heading(doc, "1.3.1 Actors", 4)
    add_table(doc, ["#", "Actor", "Description"], actors)
    add_heading(doc, "1.3.2 Use Cases (UC)", 4)
    add_table(doc, ["ID", "Use Case", "Feature", "Use Case Description"], ucs)
    add_heading(doc, "1.3.3 Use Case Diagrams", 4)
    add_para(doc, "Diagram summary: Student -> UC-01..20; Staff -> UC-21..32; StaffManager -> UC-33..34; Admin -> UC-35..40.")

    add_heading(doc, "1.4 System Functionalities", 3)
    add_heading(doc, "1.4.1 Screens Flow", 4)
    add_para(doc, "Home/Login/Register -> role dashboard -> Student learning/assessment/AI/support/profile; Staff content/questions/assessments/grading/tickets/students; Manager review queue/content pipeline/deleted topics; Admin users/settings/reports.")
    add_heading(doc, "1.4.2 Screen Authorization", 4)
    add_table(doc, ["Screen", "Student", "Staff", "StaffManager", "Admin", "Notes"], screen_auth)
    add_heading(doc, "1.4.3 Non-UI Functions", 4)
    add_table(doc, ["#", "Feature", "System Function", "Description"], non_ui)

    add_heading(doc, "1.5 Entity Relationship Diagram", 3)
    add_para(doc, "Conceptual ERD: users submit attempts/submissions, courses contain lessons, assessments contain questions, attempts contain answers, progress and flashcards track learning, support/notification/settings/audit support operations.")
    add_heading(doc, "Entities Description", 4)
    add_table(doc, ["#", "Entity", "Description"], entities)

    add_heading(doc, "2. Use Case Specifications", 2)
    uc_specs = [
        ("2.1.1 UC-01 User Login", "Student, Staff, StaffManager, Admin", "Google OAuth Provider", "Account exists and is active.", "JWT/refresh token issued and login tracked."),
        ("2.1.2 UC-02 User Register", "Student", "SMTP Server", "Email is not registered.", "Student account is verified and active."),
        ("2.2.1 UC-10 Take JLPT Mock Test", "Student", "None", "Exam is published and user has level/subscription access.", "New immutable attempt with valid score."),
        ("2.3.1 UC-13 Speaking Practice & AI Grading", "Student", "AI Speech Service, Staff", "Audio is valid and lesson is published.", "Submission/job stored; AI suggestion available or fallback returned."),
        ("2.4.1 UC-33 Review Submitted Content", "StaffManager", "Staff", "Content is pending_review.", "Content status updated and audit logged."),
    ]
    for title_text, primary, secondary, pre, post in uc_specs:
        add_heading(doc, title_text, 3)
        add_table(doc, ["Primary Actors", primary, "Secondary Actors", secondary], [])
        add_table(doc, ["Item", "Description"], [
            ("Description", "Representative complex use case from the project workflow."),
            ("Preconditions", pre),
            ("Postconditions", post),
            ("Normal Sequence/Flow", "User performs the action, backend validates authorization and business rules, system persists the result and returns standard JSON response."),
            ("Alternative Sequences/Flows", "Invalid input, insufficient permission, expired/locked resource, external service failure, or duplicate submission."),
        ])

    add_heading(doc, "3. Functional Requirements", 2)
    for heading, desc, fields in [
        ("3.1 Authentication / Login Screen", "Login/register/reset flows with JWT, OAuth, validation, role routing, and standard JSON errors.", [("Email", "Required email format"), ("Password", "Required, never logged"), ("Login button", "Calls API and shows loading/error")]),
        ("3.2 Student Learning / Lesson Detail", "Show JLPT content, VIP locks, lesson status, linked Kanji/Kana/Vocab/Grammar, and completion action.", [("JLPT level", "N5..N1"), ("Lesson status", "LOCKED/AVAILABLE/COMPLETED"), ("Complete action", "Backend updates progress")]),
        ("3.3 Assessment / Quiz and Mock Test", "Attempt UI sends answers only; backend validates time and calculates score.", [("Questions", "Published assignment questions"), ("Timer", "UX only; backend validates"), ("Submit", "Answers only, no score")]),
        ("3.4 Staff/Manager/Admin", "Separate role areas for content CRUD, review, publishing, user/settings/report management.", [("Status", "draft/pending_review/published/archived"), ("Review action", "Approve/Reject/Request Changes"), ("Audit", "actor/action/target/reason/timestamp")]),
    ]:
        add_heading(doc, heading, 3)
        add_para(doc, desc)
        add_table(doc, ["Field Name", "Description"], fields)

    add_heading(doc, "4. Non-Functional Requirements", 2)
    add_heading(doc, "4.1 External Interfaces", 3)
    add_para(doc, "REST API uses /api/[resource] and returns { status, message, data }. External integrations include MySQL 8, file storage (/uploads or S3), OCR/Speech AI, SMTP, and optional Google OAuth.")
    add_heading(doc, "4.2 Quality Attributes", 3)
    add_table(doc, ["Attribute", "Requirement"], [
        ("Usability", "Responsive Vietnamese-first web UI with clear loading, error, and empty states."),
        ("Performance", "Common API average response under 2 seconds; AI returns job_id immediately and processes in background."),
        ("Security", "JWT for private APIs, bcrypt cost >= 10, DTO-only API responses, no hardcoded secrets."),
        ("Reliability", "Soft delete, immutable submitted attempts, retry/fallback for AI/email, global exception handling."),
        ("Maintainability", "Controller -> Service -> Repository -> Entity; business logic stays in backend Service layer."),
    ])

    add_heading(doc, "5. Requirement Appendix", 2)
    add_heading(doc, "5.1 Business Rules", 3)
    add_table(doc, ["ID", "Rule Definition"], brs)
    add_heading(doc, "5.2 System Messages", 3)
    add_table(doc, ["#", "Message code", "Message Type", "Context", "Content"], messages)
    add_heading(doc, "5.3 Other Requirements", 3)
    for item in [
        "No secrets/passwords/API keys in source control.",
        "No schema change without Flyway/Liquibase migration.",
        "No frontend business logic for score, authorization, subscription, or learning progress.",
        "Reference sources: AGENTS.md, CLAUDE.md, shared_context.md, use-cases/Bao_cao_dac_ta_Use_Case.md.",
    ]:
        add_para(doc, f"- {item}")

    doc.save(path)


def main():
    outputs = {
        OUT / "SRS-Japanese-Skill-Practice-Platform.vi.md": make_markdown("vi"),
        OUT / "SRS-Japanese-Skill-Practice-Platform.en.md": make_markdown("en"),
    }
    for path, content in outputs.items():
        path.write_text(content, encoding="utf-8", newline="\n")
    make_docx("vi", OUT / "SRS-Japanese-Skill-Practice-Platform.vi.docx")
    make_docx("en", OUT / "SRS-Japanese-Skill-Practice-Platform.en.docx")


if __name__ == "__main__":
    from build_template1_srs_docs_v2 import main as main_v2

    main_v2()
