from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_template1_srs_docs as base


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "01-SRS-Requirements"
TEMPLATE = ROOT / "Guides  Templates-20260721" / "Template1_SRS Document.docx"
LOGO = ROOT / "Guides  Templates-20260721" / "images_srs" / "img_rId8.png"
DIAGRAMS = ROOT / "docs" / "07-Release-Documents" / "diagrams" / "rds"


screen_descriptions_vi = [
    ("1", "Public & Auth", "/, /login, /register, /forgot-password, /reset-password, /verify-email", "Trang vào hệ thống, xác thực, đăng ký, xác minh email và khôi phục mật khẩu."),
    ("2", "Student Dashboard", "/dashboard, /onboarding, /profile, /settings/*", "Tổng quan học tập, hồ sơ, thiết lập mục tiêu và cài đặt tài khoản."),
    ("3", "Core Learning", "/lessons/:id, /kanji, /kanji/:id, /grammar, /kana, /vocabulary", "Học Kanji, Kana, từ vựng, ngữ pháp và bài học theo cấp độ JLPT."),
    ("4", "Assessment & Review", "/quiz, /mock-test, /mock-test/:id/attempt, /mock-test/:id/results, /vocabulary/flashcard, /notebook", "Làm quiz/mock test, xem kết quả, ôn flashcard SRS và sổ tay."),
    ("5", "AI Skills", "/speaking, /kanji/:id/practice", "Luyện nói bằng audio và luyện viết tay Kanji bằng OCR/similarity."),
    ("6", "Support & Notification", "/support, /support/tickets/:id, /notifications", "Gửi ticket hỗ trợ, theo dõi phản hồi và xem thông báo."),
    ("7", "Staff Area", "/staff/*", "Quản lý nội dung, câu hỏi, bài đánh giá, chấm speaking, ticket và học viên."),
    ("8", "Manager Area", "/manager/*", "Duyệt nội dung, theo dõi pipeline, quản lý nội dung đã xóa mềm và ticket."),
    ("9", "Admin Area", "/admin/*", "Quản lý người dùng, cài đặt hệ thống, báo cáo và notification rules."),
]

screen_descriptions_en = [
    ("1", "Public & Auth", "/, /login, /register, /forgot-password, /reset-password, /verify-email", "System entry, authentication, registration, email verification, and password recovery."),
    ("2", "Student Dashboard", "/dashboard, /onboarding, /profile, /settings/*", "Learning overview, profile, target setup, and account settings."),
    ("3", "Core Learning", "/lessons/:id, /kanji, /kanji/:id, /grammar, /kana, /vocabulary", "Kanji, Kana, vocabulary, grammar, and JLPT-level lessons."),
    ("4", "Assessment & Review", "/quiz, /mock-test, /mock-test/:id/attempt, /mock-test/:id/results, /vocabulary/flashcard, /notebook", "Quizzes/mock tests, results, SRS flashcards, and notebook."),
    ("5", "AI Skills", "/speaking, /kanji/:id/practice", "Speaking audio practice and handwritten Kanji OCR/similarity practice."),
    ("6", "Support & Notification", "/support, /support/tickets/:id, /notifications", "Support ticket submission, reply tracking, and notifications."),
    ("7", "Staff Area", "/staff/*", "Content, questions, assessments, speaking grading, tickets, and student management."),
    ("8", "Manager Area", "/manager/*", "Content review, content pipeline, soft-deleted content, and tickets."),
    ("9", "Admin Area", "/admin/*", "Users, system settings, reports, and notification rules."),
]

field_specs_vi = {
    "User Register": [
        ("Full name", "Bắt buộc, tối đa 100 ký tự."),
        ("Email", "Bắt buộc, đúng định dạng email, duy nhất trong student_users."),
        ("Password", "Bắt buộc, theo policy bảo mật; lưu dạng bcrypt hash."),
        ("Confirm password", "Bắt buộc, phải khớp Password."),
        ("OTP", "6 chữ số, hết hạn sau thời gian cấu hình, dùng để xác minh email."),
    ],
    "User Login": [
        ("Email", "Bắt buộc, đúng định dạng email."),
        ("Password", "Bắt buộc, không log plaintext."),
        ("Remember session", "Tùy chọn UX; backend vẫn kiểm soát token/session."),
        ("Google login", "Tùy chọn OAuth nếu GOOGLE_CLIENT_ID được cấu hình."),
    ],
    "Password Reset": [
        ("Email", "Bắt buộc, nhận link/token reset."),
        ("Reset token", "Token một lần, có thời hạn và bị vô hiệu hóa sau khi dùng."),
        ("New password", "Theo policy bảo mật."),
        ("Confirm new password", "Phải khớp New password."),
    ],
    "Lesson Detail": [
        ("JLPT level", "N5, N4, N3, N2, N1; backend kiểm tra subscription/level."),
        ("Lesson status", "LOCKED, AVAILABLE, COMPLETED."),
        ("Content", "Text, ví dụ, audio_url/image_url, câu hỏi liên quan."),
        ("Complete action", "Gửi event hoàn thành; backend cập nhật progress."),
    ],
    "Quiz / Mock Test Attempt": [
        ("Assessment", "Quiz hoặc exam ở trạng thái published."),
        ("Question list", "Câu hỏi từ question_assignments đã publish."),
        ("Timer", "Hiển thị UX; backend validate thời gian server-side."),
        ("Answer payload", "Danh sách câu trả lời; không bao gồm score."),
        ("Result", "Score, max_score, đúng/sai, submitted_at do backend trả."),
    ],
    "AI Practice": [
        ("File/canvas input", "Ảnh hoặc audio hợp lệ, giới hạn kích thước theo cấu hình."),
        ("job_id", "Mã job AI bất đồng bộ để poll kết quả."),
        ("AI status", "PENDING, PROCESSING, DONE, FAILED."),
        ("AI result", "OCR similarity % hoặc speech score suggestion."),
        ("Final score", "Điểm cuối do Staff/logic được phép xác nhận."),
    ],
    "Staff Content Management": [
        ("Content type", "Course, Lesson, Grammar, Vocabulary, Kanji, Question, Assessment."),
        ("Status", "draft, pending_review, published, archived."),
        ("JLPT level", "N5..N1, bắt buộc không trộn level."),
        ("Submit for review", "Chỉ Staff/Admin có quyền; ghi audit khi cần."),
    ],
    "Manager Review Queue": [
        ("Filter", "Loại nội dung, JLPT level, người gửi, trạng thái, ngày gửi."),
        ("Review action", "Approve, Reject, Request Changes."),
        ("Reason/feedback", "Bắt buộc khi reject/request changes hoặc unpublish/archive."),
        ("Audit data", "actor, action, target, reason, timestamp."),
    ],
    "Admin User Management": [
        ("Role", "STUDENT, STAFF, STAFF_MANAGER, ADMIN."),
        ("Status", "ACTIVE, SUSPENDED/INACTIVE, SOFT_DELETED."),
        ("Subscription", "FREE/VIP và thời hạn hiệu lực."),
        ("Reset password", "Tạo token/tạm mật khẩu theo policy; ghi audit."),
    ],
}

field_specs_en = {
    "User Register": [
        ("Full name", "Required, maximum 100 characters."),
        ("Email", "Required, valid email format, unique in student_users."),
        ("Password", "Required, follows security policy; stored as bcrypt hash."),
        ("Confirm password", "Required, must match Password."),
        ("OTP", "6 digits, expires by configuration, used for email verification."),
    ],
    "User Login": [
        ("Email", "Required, valid email format."),
        ("Password", "Required, plaintext is never logged."),
        ("Remember session", "UX option; backend still controls token/session."),
        ("Google login", "Optional OAuth when GOOGLE_CLIENT_ID is configured."),
    ],
    "Password Reset": [
        ("Email", "Required, receives reset link/token."),
        ("Reset token", "Single-use token with expiry, invalidated after use."),
        ("New password", "Follows security policy."),
        ("Confirm new password", "Must match New password."),
    ],
    "Lesson Detail": [
        ("JLPT level", "N5, N4, N3, N2, N1; backend checks subscription/level."),
        ("Lesson status", "LOCKED, AVAILABLE, COMPLETED."),
        ("Content", "Text, examples, audio_url/image_url, linked questions."),
        ("Complete action", "Sends completion event; backend updates progress."),
    ],
    "Quiz / Mock Test Attempt": [
        ("Assessment", "Published quiz or exam."),
        ("Question list", "Questions from published question_assignments."),
        ("Timer", "UX display; backend validates server-side time."),
        ("Answer payload", "List of answers; does not include score."),
        ("Result", "Score, max_score, correctness, submitted_at returned by backend."),
    ],
    "AI Practice": [
        ("File/canvas input", "Valid image or audio within configured size limits."),
        ("job_id", "Asynchronous AI job identifier for result polling."),
        ("AI status", "PENDING, PROCESSING, DONE, FAILED."),
        ("AI result", "OCR similarity % or speech score suggestion."),
        ("Final score", "Final score confirmed by authorized Staff/logic."),
    ],
    "Staff Content Management": [
        ("Content type", "Course, Lesson, Grammar, Vocabulary, Kanji, Question, Assessment."),
        ("Status", "draft, pending_review, published, archived."),
        ("JLPT level", "N5..N1; level mixing is forbidden."),
        ("Submit for review", "Staff/Admin only; audit when required."),
    ],
    "Manager Review Queue": [
        ("Filter", "Content type, JLPT level, submitter, status, submitted date."),
        ("Review action", "Approve, Reject, Request Changes."),
        ("Reason/feedback", "Required for reject/request changes or unpublish/archive."),
        ("Audit data", "actor, action, target, reason, timestamp."),
    ],
    "Admin User Management": [
        ("Role", "STUDENT, STAFF, STAFF_MANAGER, ADMIN."),
        ("Status", "ACTIVE, SUSPENDED/INACTIVE, SOFT_DELETED."),
        ("Subscription", "FREE/VIP and validity period."),
        ("Reset password", "Creates token/temporary password by policy; audit is logged."),
    ],
}

screen_images = {
    # key -> (relative path under DIAGRAMS, English caption, Vietnamese caption)
    "User Register": ("screenshots/screen-register.png", "Real screenshot of the running app: /register", "Ảnh chụp thật từ ứng dụng đang chạy: /register"),
    "User Login": ("screenshots/screen-login.png", "Real screenshot of the running app: /login", "Ảnh chụp thật từ ứng dụng đang chạy: /login"),
    "Password Reset": ("screenshots/screen-forgot-password.png", "Real screenshot of the running app: /forgot-password", "Ảnh chụp thật từ ứng dụng đang chạy: /forgot-password"),
    "Lesson Detail": ("wf-lesson-detail.png", "Wireframe mockup (backend/DB unavailable in the build environment for a live screenshot).", "Wireframe mockup (môi trường build không có backend/DB để chụp ảnh thật)."),
    "Quiz / Mock Test Attempt": ("wf-quiz-attempt.png", "Wireframe mockup (backend/DB unavailable in the build environment for a live screenshot).", "Wireframe mockup (môi trường build không có backend/DB để chụp ảnh thật)."),
    "AI Practice": ("wf-ai-practice.png", "Wireframe mockup (backend/DB unavailable in the build environment for a live screenshot).", "Wireframe mockup (môi trường build không có backend/DB để chụp ảnh thật)."),
    "Staff Content Management": ("wf-staff-content.png", "Wireframe mockup (backend/DB unavailable in the build environment for a live screenshot).", "Wireframe mockup (môi trường build không có backend/DB để chụp ảnh thật)."),
    "Manager Review Queue": ("wf-manager-review-queue.png", "Wireframe mockup (backend/DB unavailable in the build environment for a live screenshot).", "Wireframe mockup (môi trường build không có backend/DB để chụp ảnh thật)."),
    "Admin User Management": ("wf-admin-user-management.png", "Wireframe mockup (backend/DB unavailable in the build environment for a live screenshot).", "Wireframe mockup (môi trường build không có backend/DB để chụp ảnh thật)."),
}


def _flow(*steps):
    return "\n".join(f"{i}. {s}" for i, s in enumerate(steps, 1))


def _alts(*items):
    return "\n".join(items)


uc_specs_vi = [
    ("2.1 Authentication", "2.1.1 UC-01 User Login", "Student, Staff, StaffManager, Admin", "Google OAuth Provider",
     "Người dùng mở trang đăng nhập, nhập email và mật khẩu (hoặc bấm Đăng nhập bằng Google) để vào hệ thống. Sau khi đăng nhập thành công, người dùng được vào đúng khu vực chức năng theo vai trò của mình.",
     "Tài khoản đã tồn tại, đang hoạt động (active) và không bị xóa mềm.",
     "Người dùng đăng nhập thành công, hệ thống lưu lại lượt đăng nhập, và người dùng được đưa tới đúng trang chủ theo vai trò.",
     _flow(
         "Người dùng mở trang Đăng nhập.",
         "Người dùng nhập email và mật khẩu.",
         "Người dùng bấm nút Đăng nhập.",
         "Hệ thống kiểm tra email và mật khẩu có đúng không.",
         "Hệ thống kiểm tra tài khoản còn hoạt động, chưa bị khóa.",
         "Hệ thống tạo JWT token và refresh token.",
         "Hệ thống lưu lại lượt đăng nhập vào nhật ký (audit log).",
         "Hệ thống đưa người dùng tới đúng trang chủ theo vai trò (Student/Staff/StaffManager/Admin).",
     ),
     _alts(
         "Nhánh 1 - Đăng nhập bằng Google: Người dùng bấm 'Đăng nhập bằng Google' thay vì nhập mật khẩu. Hệ thống nhờ Google xác minh người dùng, nếu đúng thì quay lại bước 6 của luồng chính.",
         "Nhánh 2 - Sai email hoặc mật khẩu: Hệ thống hiện thông báo 'Email hoặc mật khẩu không đúng' và cho phép thử lại.",
         "Nhánh 3 - Email chưa xác minh: Hệ thống yêu cầu người dùng xác minh email trước khi đăng nhập.",
         "Nhánh 4 - Tài khoản bị khóa: Nếu nhập sai mật khẩu 5 lần liên tiếp, hệ thống khóa tài khoản 15 phút và hiện cảnh báo.",
         "Nhánh 5 - Không đủ quyền truy cập trang: Nếu người dùng đã đăng nhập nhưng mở một trang không thuộc vai trò của mình, hệ thống hiện 'Bạn không có quyền truy cập'.",
     )),
    ("2.1 Authentication", "2.1.2 UC-02 User Register", "Student", "SMTP Server",
     "Một người dùng mới (Guest) tạo tài khoản Student bằng cách nhập họ tên, email và mật khẩu. Hệ thống kiểm tra thông tin rồi gửi một mã OTP tới email để xác minh email đó là thật.",
     "Email chưa được dùng để đăng ký tài khoản nào khác.",
     "Một tài khoản Student mới được tạo. Tài khoản ở trạng thái chưa xác minh cho tới khi người dùng nhập đúng mã OTP; sau đó tài khoản chuyển sang trạng thái hoạt động.",
     _flow(
         "Người dùng mở trang Đăng ký.",
         "Người dùng nhập họ tên, email, mật khẩu và xác nhận mật khẩu.",
         "Người dùng bấm nút Đăng ký.",
         "Hệ thống kiểm tra tất cả các trường hợp lệ (đúng định dạng email, mật khẩu đủ mạnh, hai mật khẩu khớp nhau).",
         "Hệ thống kiểm tra email chưa từng được dùng.",
         "Hệ thống tạo tài khoản mới với trạng thái 'chưa xác minh'.",
         "Hệ thống tạo mã OTP gồm 6 chữ số kèm thời hạn sử dụng.",
         "Hệ thống gửi mã OTP tới email của người dùng.",
         "Người dùng nhập mã OTP vào trang Xác minh Email.",
         "Hệ thống kiểm tra mã OTP đúng và kích hoạt tài khoản.",
     ),
     _alts(
         "Nhánh 1 - Email đã được đăng ký: Hệ thống hiện 'Email này đã được sử dụng' và không tạo tài khoản mới.",
         "Nhánh 2 - Mã OTP sai hoặc đã hết hạn: Hệ thống hiện lỗi và cho phép nhập lại hoặc gửi lại mã mới.",
         "Nhánh 3 - Gửi lại OTP quá nhiều lần: Nếu người dùng xin gửi lại mã quá nhiều lần trong thời gian ngắn, hệ thống tạm chặn yêu cầu mới (rate limit).",
     )),
    ("2.1 Authentication", "2.1.3 UC-03 Reset Password", "Student, Guest", "SMTP Server",
     "Người dùng quên mật khẩu có thể yêu cầu hệ thống gửi một email chứa link đặt lại mật khẩu, sau đó chọn mật khẩu mới.",
     "Email nhập vào thuộc về một tài khoản đang tồn tại và đang hoạt động.",
     "Mật khẩu được đổi và lưu dưới dạng hash mới. Token/link reset cũ không thể dùng lại lần nữa.",
     _flow(
         "Người dùng mở trang Quên mật khẩu.",
         "Người dùng nhập email và bấm Gửi.",
         "Hệ thống kiểm tra email có tồn tại không.",
         "Hệ thống tạo một reset token kèm thời hạn sử dụng.",
         "Hệ thống gửi email chứa link đặt lại mật khẩu (có kèm token).",
         "Người dùng bấm vào link và mở trang Đặt lại mật khẩu.",
         "Người dùng nhập mật khẩu mới và xác nhận lại.",
         "Hệ thống kiểm tra token còn hợp lệ, chưa hết hạn.",
         "Hệ thống lưu mật khẩu mới (đã hash) và đánh dấu token đã dùng.",
     ),
     _alts(
         "Nhánh 1 - Token hết hạn hoặc đã dùng rồi: Hệ thống báo lỗi và yêu cầu người dùng xin link reset mới.",
         "Nhánh 2 - Mật khẩu mới quá yếu: Hệ thống hiện rõ quy tắc mật khẩu và yêu cầu chọn mật khẩu mạnh hơn.",
         "Nhánh 3 - Tài khoản đang bị khóa: Hệ thống vẫn cho đổi mật khẩu, nhưng báo rõ tài khoản vẫn bị khóa cho tới khi hết thời gian khóa.",
     )),
    ("2.2 Student Learning", "2.2.1 UC-06/07/08/09 Learn Content", "Student", "None",
     "Student duyệt xem nội dung học (Ngữ pháp, Kanji, Kana, hoặc Từ vựng) theo cấp độ JLPT của mình. Student có thể đọc bài học, nghe audio, xem ví dụ, rồi đánh dấu hoàn thành hoặc bookmark để xem lại sau.",
     "Student đã đăng nhập và có quyền truy cập đúng cấp độ JLPT (khớp với subscription/level của mình).",
     "Nếu Student hoàn thành bài học, tiến độ học tập được cập nhật (chỉ tăng, không bao giờ giảm). Nếu Student bookmark, nội dung được lưu vào danh sách của Student.",
     _flow(
         "Student mở menu học tập (Ngữ pháp, Kanji, Kana, hoặc Từ vựng).",
         "Student chọn cấp độ JLPT và chủ đề.",
         "Hệ thống kiểm tra subscription và level của Student.",
         "Hệ thống kiểm tra bài học trước đó (theo thứ tự) đã hoàn thành chưa.",
         "Hệ thống hiển thị nội dung bài học (chữ, audio, hình ảnh, ví dụ).",
         "Student đọc/nghe nội dung.",
         "Student bấm 'Đánh dấu hoàn thành' hoặc 'Bookmark'.",
         "Hệ thống lưu tiến độ hoặc bookmark của Student.",
     ),
     _alts(
         "Nhánh 1 - Nội dung VIP, tài khoản Free: Hệ thống hiện thông báo cần nâng cấp VIP mới xem được, kèm link nâng cấp.",
         "Nhánh 2 - Bài học trước chưa hoàn thành: Hệ thống hiện bài học đang bị khóa và giải thích cần hoàn thành bài nào trước.",
         "Nhánh 3 - Nội dung chưa được publish: Hệ thống không hiển thị bài học này (Staff/Manager chưa duyệt/publish).",
     )),
    ("2.2 Student Learning", "2.2.2 UC-12 Flashcard Learning", "Student", "None",
     "Student ôn tập bằng flashcard theo phương pháp lặp lại ngắt quãng (SRS). Hệ thống chọn ra các thẻ đến hạn ôn, Student trả lời hoặc tự đánh giá mức độ nhớ, rồi hệ thống tính ngày ôn tiếp theo.",
     "Student có ít nhất một bộ flashcard (deck) còn hoạt động, chưa bị xóa.",
     "Ngày ôn tiếp theo (next_review_at) và hệ số dễ nhớ (ease factor) của từng thẻ đã ôn được cập nhật theo câu trả lời của Student.",
     _flow(
         "Student mở trang Flashcard/Sổ tay.",
         "Hệ thống tìm các thẻ đến hạn ôn hôm nay.",
         "Hệ thống hiện từng thẻ một (mặt câu hỏi trước).",
         "Student cố nhớ câu trả lời, rồi lật thẻ để kiểm tra.",
         "Student tự đánh giá (ví dụ: Quên, Khó, Tốt, Dễ).",
         "Hệ thống tính ngày ôn tiếp theo theo công thức SRS.",
         "Hệ thống lưu lịch ôn mới và chuyển sang thẻ kế tiếp.",
     ),
     _alts(
         "Nhánh 1 - Không có thẻ nào đến hạn: Hệ thống hiện thông báo thân thiện như 'Hôm nay không có thẻ cần ôn' và gợi ý học nội dung mới.",
         "Nhánh 2 - Deck đã bị xóa: Hệ thống báo cho Student biết bộ thẻ này không còn khả dụng.",
     )),
    ("2.3 Assessment", "2.3.1 UC-10 Take JLPT Mock Test", "Student", "None",
     "Student làm một đề thi thử JLPT đầy đủ, có tính giờ. Hệ thống theo dõi thời gian trên server, thu thập câu trả lời của Student, và tính điểm sau khi Student nộp bài hoặc khi hết giờ.",
     "Đề thi đã được publish và phù hợp với cấp độ JLPT/subscription của Student.",
     "Một bản ghi attempt mới được lưu kèm điểm số, và bản ghi này không thể bị sửa sau đó (bất biến).",
     _flow(
         "Student mở danh sách Mock Test và chọn một đề.",
         "Hệ thống kiểm tra Student có đủ quyền truy cập (level/subscription).",
         "Hệ thống tạo một attempt mới và bắt đầu đếm giờ phía server.",
         "Student trả lời câu hỏi lần lượt theo từng phần.",
         "Student bấm Nộp bài (hoặc hết giờ).",
         "Hệ thống kiểm tra thời gian thực tế đã dùng trên server (không dựa vào đồng hồ máy Student).",
         "Hệ thống tính điểm từng phần và tổng điểm.",
         "Hệ thống lưu attempt ở trạng thái bất biến (không thể sửa sau này).",
         "Hệ thống hiển thị màn hình kết quả cho Student.",
     ),
     _alts(
         "Nhánh 1 - Hết giờ: Hệ thống tự động nộp các câu Student đã chọn tới thời điểm đó.",
         "Nhánh 2 - Student cố nộp bài 2 lần: Hệ thống chặn lần nộp thứ hai và hiện lại kết quả đã có.",
         "Nhánh 3 - Đề bị unpublish trong lúc đang làm: Hệ thống vẫn cho Student hoàn thành attempt đang làm dở, nhưng đề sẽ không xuất hiện cho lượt làm mới.",
         "Nhánh 4 - Thiếu quyền VIP: Hệ thống chặn Student bắt đầu làm bài và hiện thông báo nâng cấp.",
     )),
    ("2.3 Assessment", "2.3.2 UC-11 Practice & Quiz", "Student", "None",
     "Student làm một bài quiz ngắn về một bài học/chủ đề, dùng để luyện tập hằng ngày, không phải thi thử đầy đủ. Hệ thống chấm điểm ngay sau khi nộp và hiện kết quả.",
     "Quiz đã được publish. Các câu hỏi trong quiz đúng cấp độ JLPT của Student.",
     "Một attempt mới và kết quả được lưu và hiển thị cho Student.",
     _flow(
         "Student mở một bài học/chủ đề và bấm 'Bắt đầu Quiz'.",
         "Hệ thống gửi danh sách câu hỏi (không kèm đáp án đúng).",
         "Student trả lời từng câu.",
         "Student bấm Nộp bài.",
         "Hệ thống kiểm tra đáp án và tính điểm trên server.",
         "Hệ thống lưu attempt mới.",
         "Hệ thống hiện điểm số và câu nào đúng/sai.",
     ),
     _alts(
         "Nhánh 1 - Quiz đã bị archive: Hệ thống không cho làm bài mới và hiện thông báo.",
         "Nhánh 2 - Câu hỏi bị khóa sau lần làm đầu tiên: Nếu Staff cố sửa một quiz đã có người làm, hệ thống chặn sửa và yêu cầu tạo phiên bản mới.",
         "Nhánh 3 - Dữ liệu nhập không hợp lệ: Nếu thiếu câu trả lời hoặc sai định dạng, hệ thống báo lỗi và không cho nộp bài.",
     )),
    ("2.4 AI Skills", "2.4.1 UC-13 Speaking Practice & AI Grading", "Student", "AI Speech Service, Staff",
     "Student ghi âm hoặc tải lên một file audio trả lời cho bài luyện nói. Hệ thống gửi audio cho dịch vụ AI xử lý ở chế độ nền, sau đó hiện điểm AI gợi ý. Một Staff có thể xem lại và xác nhận điểm cuối cùng.",
     "File audio hợp lệ (đúng định dạng/kích thước). Bài luyện nói đã được publish.",
     "Một bản ghi submission được lưu. Trạng thái AI luôn rõ ràng (PENDING, PROCESSING, DONE, hoặc FAILED) để Student luôn biết chuyện gì đang xảy ra.",
     _flow(
         "Student ghi âm hoặc tải lên file audio cho bài luyện nói.",
         "Student bấm Nộp bài.",
         "Hệ thống lưu file audio và tạo một job với trạng thái PENDING.",
         "Hệ thống trả lời ngay cho Student kèm job_id (không bắt Student chờ).",
         "Ở chế độ nền, hệ thống gửi audio cho dịch vụ AI chấm giọng nói.",
         "Dịch vụ AI trả về điểm gợi ý.",
         "Hệ thống cập nhật trạng thái job thành DONE và lưu điểm AI gợi ý.",
         "Student kiểm tra kết quả bằng cách tải lại trang (polling).",
         "Một Staff có thể mở bài nộp, nghe audio, rồi xác nhận hoặc sửa điểm cuối cùng.",
     ),
     _alts(
         "Nhánh 1 - Dịch vụ AI chậm hoặc lỗi: Hệ thống chờ tới một mốc thời gian (timeout) rồi thử lại (tối đa 3 lần). Nếu vẫn lỗi, trạng thái job chuyển thành FAILED và Student thấy thông báo rõ ràng, không phải màn hình trắng.",
         "Nhánh 2 - Sai định dạng file: Hệ thống từ chối file trước khi upload và giải thích các định dạng được phép.",
         "Nhánh 3 - Staff sửa điểm AI: Staff có thể ghi đè điểm AI gợi ý bằng điểm cuối cùng của mình kèm nhận xét.",
     )),
    ("2.4 AI Skills", "2.4.2 UC-20 AI Handwriting Practice", "Student", "OCR AI Service",
     "Student viết tay một chữ Kanji (trên canvas hoặc tải ảnh lên) và nhờ hệ thống kiểm tra xem gần giống chữ đúng tới mức nào. AI chỉ so sánh hình dạng nét vẽ, không kiểm tra thứ tự nét.",
     "Ảnh/canvas vẽ hợp lệ. Chữ Kanji mục tiêu tồn tại trong hệ thống.",
     "Một tỉ lệ giống nhau (similarity %) được lưu và hiển thị cho Student.",
     _flow(
         "Student mở một chữ Kanji và bấm 'Luyện viết'.",
         "Student viết chữ Kanji trên canvas (hoặc tải ảnh lên).",
         "Student bấm Nộp bài.",
         "Hệ thống lưu ảnh và tạo một OCR job.",
         "Hệ thống gửi ảnh cho dịch vụ AI OCR.",
         "Dịch vụ AI so sánh hình dạng với chữ đúng và trả về tỉ lệ giống nhau.",
         "Hệ thống hiện kết quả kèm nhận xét đơn giản (ví dụ: 'Khá tốt' hoặc 'Cần luyện thêm').",
     ),
     _alts(
         "Nhánh 1 - Dịch vụ AI lỗi: Hệ thống hiện lỗi rõ ràng và cho phép Student thử lại.",
         "Nhánh 2 - Ảnh quá lớn: Hệ thống từ chối file và báo giới hạn kích thước.",
         "Nhánh 3 - Không kiểm tra thứ tự nét: Đây là hành vi có chủ đích, không phải lỗi — hệ thống chỉ kiểm tra hình dạng cuối cùng.",
     )),
    ("2.5 Staff Content Management", "2.5.1 UC-24/25/26/27/28 Manage Content", "Staff", "StaffManager",
     "Một Staff tạo hoặc sửa nội dung học: bài học, ngữ pháp, từ vựng, Kanji, câu hỏi, quiz, hoặc đề thi thử. Khi nội dung đã sẵn sàng, Staff gửi cho StaffManager duyệt trước khi hiển thị cho Student.",
     "Tài khoản Staff đang hoạt động và có quyền quản lý loại nội dung này.",
     "Nội dung được lưu ở trạng thái draft (đang soạn) hoặc pending_review (đã gửi duyệt). Nếu là thao tác quan trọng, hành động được ghi vào audit log.",
     _flow(
         "Staff mở màn hình quản lý nội dung (Bài học, Câu hỏi, hoặc Đề thi).",
         "Staff tạo mới hoặc mở một mục có sẵn để sửa.",
         "Staff nhập các trường nội dung (tiêu đề, cấp độ JLPT, nội dung, ví dụ...).",
         "Hệ thống kiểm tra các trường theo quy tắc validate và business rule.",
         "Staff bấm 'Lưu nháp' hoặc 'Gửi duyệt'.",
         "Hệ thống lưu nội dung ở trạng thái draft hoặc pending_review.",
     ),
     _alts(
         "Nhánh 1 - Nội dung đã có Student làm bài: Hệ thống khóa việc sửa phiên bản đó và yêu cầu Staff tạo phiên bản mới.",
         "Nhánh 2 - Thiếu dữ liệu bắt buộc: Hệ thống đánh dấu các trường còn thiếu và không cho lưu cho tới khi bổ sung đủ.",
         "Nhánh 3 - Sai cấp độ JLPT: Hệ thống không cho phép trộn nội dung của nhiều cấp độ khác nhau trong cùng một bài học.",
     )),
    ("2.6 Content Review", "2.6.1 UC-33/34 Review and Publish Content", "StaffManager", "Staff",
     "Một StaffManager kiểm tra nội dung mà Staff đã gửi duyệt. StaffManager có thể duyệt (Approve), từ chối (Reject), hoặc yêu cầu sửa lại (Request Changes). Sau khi được duyệt, StaffManager có thể publish để Student xem được, hoặc sau này unpublish/archive/khôi phục.",
     "Nội dung đang ở trạng thái pending_review (để duyệt/từ chối/yêu cầu sửa), hoặc đã published (để unpublish/archive/khôi phục).",
     "Trạng thái nội dung được cập nhật. Hệ thống ghi audit log và gửi thông báo cho Staff đã gửi nội dung đó.",
     _flow(
         "StaffManager mở hàng đợi duyệt (Review Queue).",
         "StaffManager mở một mục nội dung để xem chi tiết đầy đủ.",
         "StaffManager chọn hành động: Duyệt, Từ chối, hoặc Yêu cầu sửa.",
         "Nếu là Từ chối hoặc Yêu cầu sửa, StaffManager nhập lý do.",
         "Hệ thống cập nhật trạng thái nội dung và lưu audit log.",
         "Hệ thống gửi thông báo cho Staff về quyết định này.",
         "Với nội dung đã duyệt, StaffManager có thể bấm Publish để hiển thị cho Student.",
     ),
     _alts(
         "Nhánh 1 - Nội dung đã được xử lý rồi: Nếu một reviewer khác đã duyệt/từ chối nội dung này trước, hệ thống hiện thông báo và làm mới danh sách.",
         "Nhánh 2 - Không đủ quyền: Nếu người dùng là Staff nhưng không phải StaffManager, hệ thống chặn hành động với lỗi 403.",
         "Nhánh 3 - Thiếu lý do: Hệ thống bắt buộc nhập lý do cho các hành động Từ chối, Yêu cầu sửa, Unpublish, và Archive.",
     )),
    ("2.7 Administration", "2.7.1 UC-37/39/40 Admin Management", "Admin", "SMTP Server",
     "Một Admin quản lý tài khoản người dùng, cấu hình hệ thống, và các quy tắc thông báo tự động. Đây là các thao tác quản trị hằng ngày như đổi vai trò người dùng, cập nhật cấu hình SMTP, hoặc tạo một quy tắc thông báo mới.",
     "Admin đã đăng nhập bằng một tài khoản Admin hợp lệ.",
     "Dữ liệu đã thay đổi (user, setting, hoặc rule) được lưu lại, và thay đổi được ghi vào audit log.",
     _flow(
         "Admin mở khu vực quản trị (Users, Settings, hoặc Notification Rules).",
         "Admin tìm kiếm hoặc lọc để tìm đúng bản ghi cần sửa.",
         "Admin mở bản ghi và cập nhật thông tin.",
         "Hệ thống kiểm tra dữ liệu mới theo quy tắc validate.",
         "Admin bấm Lưu.",
         "Hệ thống lưu thay đổi và ghi một dòng audit log.",
     ),
     _alts(
         "Nhánh 1 - Hạ quyền Admin cuối cùng: Hệ thống chặn hành động này để hệ thống luôn còn ít nhất một tài khoản Admin.",
         "Nhánh 2 - Giá trị setting không hợp lệ: Hệ thống báo lỗi và không lưu setting đó.",
         "Nhánh 3 - Quy tắc thông báo có thể gây spam: Hệ thống cảnh báo Admin nếu một rule sẽ gửi quá nhiều thông báo trong thời gian ngắn.",
     )),
]

uc_specs_en = [
    ("2.1 Authentication", "2.1.1 UC-01 User Login", "Student, Staff, StaffManager, Admin", "Google OAuth Provider",
     "A user opens the login page and enters their email and password (or clicks Login with Google) to sign in. After a successful login, the system sends the user to the correct area for their role.",
     "The user already has an account. The account is active and is not soft-deleted.",
     "The user is logged in, the system saves a login record, and the user sees the correct home page for their role.",
     _flow(
         "The user opens the Login page.",
         "The user types their email and password.",
         "The user clicks the Login button.",
         "The system checks the email and password.",
         "The system checks that the account is active and not locked.",
         "The system creates a JWT token and a refresh token.",
         "The system saves a login record (audit log).",
         "The system sends the user to the correct home page based on their role (Student, Staff, StaffManager, or Admin).",
     ),
     _alts(
         "Alt 1 - Login with Google: The user clicks 'Login with Google' instead of typing a password. The system asks Google to check the user. If Google confirms the user, the system goes back to step 6 of the Normal Flow.",
         "Alt 2 - Wrong email or password: The system shows the message 'Email or password is incorrect' and lets the user try again.",
         "Alt 3 - Email not verified: The system tells the user to verify their email before they can log in.",
         "Alt 4 - Account locked: If the user enters the wrong password 5 times in a row, the system locks the account for 15 minutes and shows a warning.",
         "Alt 5 - No permission for a page: If a logged-in user opens a page that does not belong to their role, the system shows 'You do not have permission to access this page'.",
     )),
    ("2.1 Authentication", "2.1.2 UC-02 User Register", "Student", "SMTP Server",
     "A new user (Guest) creates a Student account by filling in their name, email, and password. The system checks the information and sends a one-time code (OTP) to the user's email to prove the email is real.",
     "The email address is not used by another account yet.",
     "A new Student account is created. The account stays inactive until the user enters the correct OTP; after that, the account becomes active.",
     _flow(
         "The user opens the Register page.",
         "The user types their full name, email, password, and confirms the password.",
         "The user clicks the Register button.",
         "The system checks that all fields are valid (for example: email format, password strength, passwords match).",
         "The system checks that the email is not already used.",
         "The system creates the new account with status 'not verified'.",
         "The system creates a 6-digit OTP code with an expiry time.",
         "The system sends the OTP code to the user's email.",
         "The user enters the OTP code on the Verify Email page.",
         "The system checks the OTP code and activates the account.",
     ),
     _alts(
         "Alt 1 - Email already used: The system shows 'This email is already registered' and does not create a new account.",
         "Alt 2 - Wrong or expired OTP: The system shows an error and lets the user try again or ask for a new code.",
         "Alt 3 - Too many resend requests: If the user asks for a new OTP too many times in a short time, the system blocks new requests for a while (rate limit).",
     )),
    ("2.1 Authentication", "2.1.3 UC-03 Reset Password", "Student, Guest", "SMTP Server",
     "A user who forgot their password asks the system to send a password reset link by email, then chooses a new password.",
     "The email belongs to an existing, active account.",
     "The password is changed and saved as a new hash. The old reset link/token can no longer be used.",
     _flow(
         "The user opens the Forgot Password page.",
         "The user types their email and clicks Send.",
         "The system checks that the email exists.",
         "The system creates a reset token with an expiry time.",
         "The system sends a reset email with a link that contains the token.",
         "The user clicks the link and opens the Reset Password page.",
         "The user types a new password and confirms it.",
         "The system checks that the token is valid and not expired.",
         "The system saves the new password (hashed) and marks the token as used.",
     ),
     _alts(
         "Alt 1 - Expired or already-used token: The system shows an error and asks the user to request a new reset email.",
         "Alt 2 - Weak new password: The system shows the password rules and asks the user to choose a stronger password.",
         "Alt 3 - Account is locked: The system still lets the user reset the password, but tells them the account stays locked until the lock time ends.",
     )),
    ("2.2 Student Learning", "2.2.1 UC-06/07/08/09 Learn Content", "Student", "None",
     "A Student browses learning content (Grammar, Kanji, Kana, or Vocabulary) for their JLPT level. The Student can read a lesson, listen to audio, look at examples, and mark it as complete or bookmark it for later.",
     "The Student is signed in and has access to the chosen JLPT level (matches their subscription/level).",
     "If the Student completes the lesson, their progress is updated (it only goes up, never down). If the Student bookmarks the content, it is saved to their list.",
     _flow(
         "The Student opens the learning menu (Grammar, Kanji, Kana, or Vocabulary).",
         "The Student picks a JLPT level and a topic.",
         "The system checks the Student's subscription and level.",
         "The system checks that the previous lesson in order is already completed.",
         "The system shows the lesson content (text, audio, image, examples).",
         "The Student reads or listens to the content.",
         "The Student clicks 'Mark as Complete' or 'Bookmark'.",
         "The system saves the Student's progress or bookmark.",
     ),
     _alts(
         "Alt 1 - VIP content, Free account: The system shows a message that this content needs a VIP subscription and offers an upgrade link.",
         "Alt 2 - Previous lesson not finished: The system shows the lesson as locked and explains which lesson must be finished first.",
         "Alt 3 - Content not published yet: The system does not show the lesson, because Staff/Manager has not published it.",
     )),
    ("2.2 Student Learning", "2.2.2 UC-12 Flashcard Learning", "Student", "None",
     "A Student reviews flashcards using Spaced Repetition (SRS). The system picks the cards that are due for review, the Student answers or rates how well they remembered each card, and the system plans the next review date.",
     "The Student has at least one flashcard deck with cards that are not deleted.",
     "The next review date and the ease factor of each reviewed card are updated based on the Student's answers.",
     _flow(
         "The Student opens the Flashcard/Notebook page.",
         "The system finds the cards that are due for review today.",
         "The system shows one card at a time (question side first).",
         "The Student tries to recall the answer, then flips the card to check.",
         "The Student rates the answer (for example: Again, Hard, Good, Easy).",
         "The system calculates the next review date using the SRS rule.",
         "The system saves the updated schedule and moves to the next card.",
     ),
     _alts(
         "Alt 1 - No cards are due: The system shows a friendly message like 'No cards to review today' and suggests learning new content instead.",
         "Alt 2 - Deck was deleted: The system tells the Student this deck is no longer available.",
     )),
    ("2.3 Assessment", "2.3.1 UC-10 Take JLPT Mock Test", "Student", "None",
     "A Student takes a full-length, timed JLPT mock exam. The system tracks the time on the server, collects the Student's answers, and calculates the score after the Student submits or when time runs out.",
     "The mock exam is published and matches the Student's JLPT level/subscription.",
     "A new attempt record is saved with the score, and this record can never be changed afterward (immutable).",
     _flow(
         "The Student opens the Mock Test list and picks an exam.",
         "The system checks that the Student has access (level/subscription).",
         "The system creates a new attempt and starts the server-side timer.",
         "The Student answers the questions, one section at a time.",
         "The Student clicks Submit (or the timer reaches zero).",
         "The system checks the real time used on the server, not the Student's device clock.",
         "The system calculates the score for each part and the total score.",
         "The system saves the attempt as immutable (it cannot be edited later).",
         "The system shows the result screen to the Student.",
     ),
     _alts(
         "Alt 1 - Time runs out: The system automatically submits the answers the Student has chosen so far.",
         "Alt 2 - The Student tries to submit twice: The system blocks the second submit and shows the existing result.",
         "Alt 3 - The exam is unpublished during the attempt: The system still lets the Student finish the attempt already in progress, but the exam will not appear for new attempts.",
         "Alt 4 - Missing VIP access: The system blocks the Student from starting and shows an upgrade message.",
     )),
    ("2.3 Assessment", "2.3.2 UC-11 Practice & Quiz", "Student", "None",
     "A Student takes a short quiz about one lesson or topic, for daily practice, not a full mock exam. The system grades the quiz right after submission and shows the result.",
     "The quiz is published. The questions in the quiz match the Student's JLPT level.",
     "A new attempt and result are saved and shown to the Student.",
     _flow(
         "The Student opens a lesson/topic and clicks 'Start Quiz'.",
         "The system sends the list of questions, without the correct answers.",
         "The Student answers each question.",
         "The Student clicks Submit.",
         "The system checks the answers and calculates the score on the server.",
         "The system saves the new attempt.",
         "The system shows the score and which answers were right or wrong.",
     ),
     _alts(
         "Alt 1 - Quiz is archived: The system does not allow a new attempt and shows a message.",
         "Alt 2 - Questions are locked after the first attempt: If Staff tries to edit a quiz that already has attempts, the system blocks the edit and asks Staff to create a new version instead.",
         "Alt 3 - Invalid input: If an answer is missing or in the wrong format, the system shows an error and does not submit.",
     )),
    ("2.4 AI Skills", "2.4.1 UC-13 Speaking Practice & AI Grading", "Student", "AI Speech Service, Staff",
     "A Student records or uploads an audio answer for a speaking exercise. The system sends the audio to an AI service in the background and later shows an AI suggested score. A Staff member can review and confirm the final score.",
     "The audio file is valid (right format and size). The speaking lesson is published.",
     "A submission record is saved. Its AI status is always clear (PENDING, PROCESSING, DONE, or FAILED), so the Student always knows what is happening.",
     _flow(
         "The Student records or uploads an audio file for the speaking task.",
         "The Student clicks Submit.",
         "The system saves the audio file and creates a job with status PENDING.",
         "The system replies to the Student right away with the job_id, without making the Student wait.",
         "In the background, the system sends the audio to the AI speech service.",
         "The AI service returns a suggested score.",
         "The system updates the job status to DONE and saves the AI suggested score.",
         "The Student checks the result by refreshing the page (polling).",
         "A Staff member can open the submission, listen to the audio, and confirm or change the final score.",
     ),
     _alts(
         "Alt 1 - The AI service is slow or fails: The system waits up to a timeout, then tries again (up to 3 times). If it still fails, the job status becomes FAILED and the Student sees a clear message, not a blank screen.",
         "Alt 2 - Wrong file type: The system rejects the file before upload and explains which formats are allowed.",
         "Alt 3 - Staff changes the AI score: The Staff member can override the AI suggested score with their own final score and a comment.",
     )),
    ("2.4 AI Skills", "2.4.2 UC-20 AI Handwriting Practice", "Student", "OCR AI Service",
     "A Student draws a Kanji character by hand (on a canvas or by uploading an image) and asks the system to check how close it is to the correct Kanji. The AI compares only the shape of the strokes, not the stroke order.",
     "The drawing/image is valid. The target Kanji exists in the system.",
     "A similarity percentage is saved and shown to the Student.",
     _flow(
         "The Student opens a Kanji and clicks 'Practice Writing'.",
         "The Student draws the Kanji on the canvas (or uploads an image).",
         "The Student clicks Submit.",
         "The system saves the image and creates an OCR job.",
         "The system sends the image to the OCR AI service.",
         "The AI service compares the shape to the correct Kanji and returns a similarity percentage.",
         "The system shows the result with a simple comment, for example 'Good try' or 'Needs more practice'.",
     ),
     _alts(
         "Alt 1 - The AI service fails: The system shows a clear error and lets the Student try again.",
         "Alt 2 - Image too large: The system rejects the file and tells the Student the size limit.",
         "Alt 3 - Stroke order is not checked: This is expected behavior, not an error — the system only checks the final shape.",
     )),
    ("2.5 Staff Content Management", "2.5.1 UC-24/25/26/27/28 Manage Content", "Staff", "StaffManager",
     "A Staff member creates or edits learning content: lessons, grammar points, vocabulary, Kanji, questions, quizzes, or mock exams. When the content is ready, Staff sends it to a StaffManager for review before it can be shown to Students.",
     "The Staff account is active and has permission to manage this type of content.",
     "The content is saved as draft (still editing) or pending_review (sent for approval). If it is an important action, it is written to the audit log.",
     _flow(
         "Staff opens the content management screen (Lessons, Questions, or Assessments).",
         "Staff creates a new item or opens an existing one to edit.",
         "Staff fills in the content fields, for example title, JLPT level, text, and examples.",
         "The system checks the fields against validation rules and business rules.",
         "Staff clicks 'Save as Draft' or 'Submit for Review'.",
         "The system saves the content with status draft or pending_review.",
     ),
     _alts(
         "Alt 1 - The content already has student attempts: The system locks editing of that version and tells Staff to create a new version instead.",
         "Alt 2 - Missing required data: The system highlights the missing fields and does not save until they are fixed.",
         "Alt 3 - Wrong JLPT level: The system does not allow mixing content from different levels in the same lesson.",
     )),
    ("2.6 Content Review", "2.6.1 UC-33/34 Review and Publish Content", "StaffManager", "Staff",
     "A StaffManager checks content that Staff submitted for review. The StaffManager can approve, reject, or ask for changes. After content is approved, the StaffManager can publish it so Students can see it, or later unpublish, archive, or restore it.",
     "The content is in pending_review status (for approve/reject/request changes), or already published (for unpublish/archive/restore).",
     "The content status is updated. The system writes an audit record and sends a notification to the Staff member who submitted the content.",
     _flow(
         "StaffManager opens the Review Queue.",
         "StaffManager opens one content item to see its full details.",
         "StaffManager chooses an action: Approve, Reject, or Request Changes.",
         "If the action is Reject or Request Changes, StaffManager types a reason.",
         "The system updates the content status and saves the audit log.",
         "The system sends a notification to the Staff member about the decision.",
         "For approved content, StaffManager can click Publish to make it visible to Students.",
     ),
     _alts(
         "Alt 1 - Content already handled: If another reviewer already approved or rejected this content, the system shows a message and refreshes the list.",
         "Alt 2 - Not enough permission: If the user is Staff but not StaffManager, the system blocks the action with a 403 error.",
         "Alt 3 - Missing reason: The system requires a reason for Reject, Request Changes, Unpublish, and Archive actions.",
     )),
    ("2.7 Administration", "2.7.1 UC-37/39/40 Admin Management", "Admin", "SMTP Server",
     "An Admin manages user accounts, system settings, and notification rules. This covers day-to-day system administration tasks like changing a user's role, updating SMTP settings, or creating a new automatic notification rule.",
     "The Admin is signed in with a valid Admin account.",
     "The changed data (user, setting, or rule) is saved, and the change is written to the audit log.",
     _flow(
         "Admin opens the Admin area (Users, Settings, or Notification Rules).",
         "Admin searches or filters to find the record to change.",
         "Admin opens the record and updates the information.",
         "The system checks the new data against validation rules.",
         "Admin clicks Save.",
         "The system saves the change and writes an audit log entry.",
     ),
     _alts(
         "Alt 1 - Demoting the last Admin: The system blocks this action, so the system always keeps at least one Admin account.",
         "Alt 2 - Invalid setting value: The system shows an error and does not save the setting.",
         "Alt 3 - Notification rule could cause spam: The system warns the Admin if a rule would send too many notifications in a short time.",
     )),
]


def rows_for(lang, key):
    return (field_specs_vi if lang == "vi" else field_specs_en)[key]


def clear_document(doc):
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def add_p(doc, text="", bold=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(10.5)
    return p


def add_h(doc, text, level):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(str(header))
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    return table


def add_toc_field(doc):
    """Add a Word table-of-contents field similar to the template TOC block."""
    paragraph = doc.add_paragraph()

    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-5" \\h \\z \\u'
    run._r.append(instr)

    run = paragraph.add_run()
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    fallback_lines = [
        "I. Record of Changes",
        "II. Software Requirement Specification",
        "1. Overall Requirements",
        "1.1 Context Diagram",
        "1.2 Main Business Processes",
        "1.3 User Requirements",
        "1.4 System Functionalities",
        "1.5 Entity Relationship Diagram",
        "2. Use Case Specifications",
        "3. Functional Requirements",
        "4. Non-Functional Requirements",
        "5. Requirement Appendix",
    ]
    for index, line in enumerate(fallback_lines):
        if index == 0:
            paragraph.add_run(line)
        else:
            doc.add_paragraph(line)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def add_uc_table(doc, spec):
    _, _, primary, secondary, description, pre, post, normal, alt = spec
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    rows = [
        ("Primary Actors", primary, "Secondary Actors", secondary),
        ("Description", description, "", ""),
        ("Preconditions", pre, "", ""),
        ("Postconditions", post, "", ""),
        ("Normal Sequence/Flow", normal, "", ""),
        ("Alternative Sequences/Flows", alt, "", ""),
    ]
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            table.rows[r].cells[c].text = value
    return table


def add_picture_if_exists(doc, path, width=6.2):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
    else:
        add_p(doc, f"[Diagram placeholder: {path.name}]")


def md_table(headers, rows):
    text = "| " + " | ".join(headers) + " |\n"
    text += "| " + " | ".join(["---"] * len(headers)) + " |\n"
    for row in rows:
        text += "| " + " | ".join(str(v).replace("\n", "<br>") for v in row) + " |\n"
    return text


def make_markdown(lang):
    vi = lang == "vi"
    actors = base.actors_vi if vi else base.actors_en
    ucs = base.use_cases_vi if vi else base.use_cases_en
    brs = base.business_rules_vi if vi else base.business_rules_en
    entities = base.entities_vi if vi else base.entities_en
    non_ui = base.non_ui_vi if vi else base.non_ui_en
    messages = base.messages_vi if vi else base.messages_en
    screens = screen_descriptions_vi if vi else screen_descriptions_en
    specs = uc_specs_vi if vi else uc_specs_en
    title = "ĐẶC TẢ YÊU CẦU PHẦN MỀM" if vi else "SOFTWARE REQUIREMENT SPECIFICATION"
    intro = (
        "Japanese Skill Practice Platform là hệ thống web học và luyện thi tiếng Nhật JLPT từ N5 đến N1, gồm React 18 frontend, Spring Boot 3/Java 21 backend và MySQL 8. Hệ thống hỗ trợ học Kanji, Kana, từ vựng, ngữ pháp, flashcard SRS, quiz, mock exam, AI OCR/Speech, hỗ trợ học viên, kiểm duyệt nội dung và quản trị."
        if vi
        else "Japanese Skill Practice Platform is a web-based JLPT learning and exam practice system from N5 to N1, built with React 18, Spring Boot 3/Java 21, and MySQL 8. It supports Kanji, Kana, vocabulary, grammar, SRS flashcards, quizzes, mock exams, OCR/Speech AI, student support, content review, and administration."
    )
    processes = [
        ("1.2.1", "Authentication and Account Lifecycle", "Register, verify email, login, reset/change password, logout, token/session management.", "bp-1-authentication.png"),
        ("1.2.2", "Student Learning Path", "Browse level content, unlock lessons by order, complete lessons, bookmark and track progress.", "bp-2-learning-path.png"),
        ("1.2.3", "Assessment and SRS Review", "Take quizzes/mock exams, backend grading, immutable attempts, flashcard spaced repetition.", "bp-3-assessment-srs.png"),
        ("1.2.4", "AI Practice Processing", "Upload image/audio, create asynchronous AI jobs, timeout/retry/fallback, poll result.", "bp-4-ai-processing.png"),
        ("1.2.5", "Content Governance", "Staff drafts content, StaffManager reviews and publishes, audit and notification.", "bp-5-content-governance.png"),
        ("1.2.6", "System Administration", "Manage users, settings, subscriptions, reports, and notification rules.", "bp-6-system-admin.png"),
    ]
    process_sections_md = "\n".join(
        f"#### {pid} {name}\n\n{desc}\n\n![{name}](../07-Release-Documents/diagrams/rds/{diagram})\n"
        for pid, name, desc, diagram in processes
    )
    md = f"""# {title}

**Project Name:** Japanese Skill Practice Platform  
**Version:** 1.0  
**Location/Date:** Hanoi, July 2026  
**Template:** `Guides  Templates-20260721/Template1_SRS Document.docx`

**Table of Contents**

- [I. Record of Changes](#i-record-of-changes)
- [II. Software Requirement Specification](#ii-software-requirement-specification)
- [1. Overall Requirements](#1-overall-requirements)
- [2. Use Case Specifications](#2-use-case-specifications)
- [3. Functional Requirements](#3-functional-requirements)
- [4. Non-Functional Requirements](#4-non-functional-requirements)
- [5. Requirement Appendix](#5-requirement-appendix)

# I. Record of Changes

| Date | A/M/D | In charge | Change Description |
| --- | --- | --- | --- |
| 2026-07-23 | A | AI Agent | Generated bilingual SRS according to Template1 DOCX structure. |
| 2026-07-23 | M | AI Agent | Revised to include missing Template1 sections: screen descriptions, fuller UC specifications, functional field descriptions, NFRs, appendix. |
| 2026-07-24 | M | AI Agent | Real Word Table of Contents field; split 1.2 into 6 sub-processes with swim-lane diagrams; added screen images to section 3 (real screenshots for Auth screens, wireframe mockups for screens requiring a logged-in session not reachable in the build environment). |
| 2026-07-24 | M | AI Agent | Expanded Use Case Specifications (section 2): longer step-by-step Normal Flow and named Alternative Flows for all 15 UC groups, written in simple, easy-to-understand English. |
| 2026-07-24 | M | AI Agent | Split 1.3.3 Use Case Diagrams into one diagram per actor (Guest, Student, Staff, StaffManager, Admin), matching Template1's per-actor layout; added Guest to the Actors table. |

*A - Added, M - Modified, D - Deleted*

# II. Software Requirement Specification

## 1. Overall Requirements

### 1.1 Context Diagram

{intro}

![Context Diagram](../07-Release-Documents/diagrams/rds/package-diagram.png)

### 1.2 Main Business Processes

{process_sections_md}

### 1.3 User Requirements

#### 1.3.1 Actors

{md_table(["#", "Actor", "Description"], actors)}

#### 1.3.2 Use Cases (UC)

{md_table(["ID", "Use Case", "Feature", "Use Case Description"], ucs)}

#### 1.3.3 Use Case Diagrams

##### 1.3.3.1 UCs for Guest

![UCs for Guest](../07-Release-Documents/diagrams/rds/uc-guest.png)

##### 1.3.3.2 UCs for Student

![UCs for Student](../07-Release-Documents/diagrams/rds/uc-student.png)

##### 1.3.3.3 UCs for Staff

![UCs for Staff](../07-Release-Documents/diagrams/rds/uc-staff.png)

##### 1.3.3.4 UCs for StaffManager

![UCs for StaffManager](../07-Release-Documents/diagrams/rds/uc-staffmanager.png)

##### 1.3.3.5 UCs for Admin

![UCs for Admin](../07-Release-Documents/diagrams/rds/uc-admin.png)

### 1.4 System Functionalities

#### 1.4.1 Screens Flow

![Screens Flow](../07-Release-Documents/diagrams/rds/screens-flow.png)

**Screen Descriptions**

{md_table(["#", "Feature", "Screen (Route)", "Description"], screens)}

#### 1.4.2 Screen Authorization

{md_table(["Screen", "Student", "Staff", "StaffManager", "Admin", "Notes"], base.screen_auth)}

#### 1.4.3 Non-UI Functions

{md_table(["#", "Feature", "System Function", "Description"], non_ui)}

### 1.5 Entity Relationship Diagram

![Entity Relationship Diagram](../07-Release-Documents/diagrams/rds/er-diagram.png)

**Entities Description**

{md_table(["#", "Entity", "Description"], entities)}

## 2. Use Case Specifications

"""
    current = None
    for spec in specs:
        feature, title_text, primary, secondary, desc, pre, post, normal, alt = spec
        if feature != current:
            md += f"### {feature}\n\n"
            current = feature
        md += f"#### {title_text}\n\n"
        md += md_table(
            ["Primary Actors", primary, "Secondary Actors", secondary],
            [
                ["Description", desc, "", ""],
                ["Preconditions", pre, "", ""],
                ["Postconditions", post, "", ""],
                ["Normal Sequence/Flow", normal, "", ""],
                ["Alternative Sequences/Flows", alt, "", ""],
            ],
        )
        md += "\n"
    md += "## 3. Functional Requirements\n\n"
    groups = [
        ("3.1 User Authentication", ["User Register", "User Login", "Password Reset"]),
        ("3.2 Student Learning and Assessment", ["Lesson Detail", "Quiz / Mock Test Attempt", "AI Practice"]),
        ("3.3 Staff, Manager and Admin Operations", ["Staff Content Management", "Manager Review Queue", "Admin User Management"]),
    ]
    for group, keys in groups:
        md += f"### {group}\n\n"
        for idx, key in enumerate(keys, 1):
            md += f"#### {group.split()[0]}.{idx} {key}\n\n"
            img_path, cap_en, cap_vi = screen_images[key]
            md += f"![{key}](../07-Release-Documents/diagrams/rds/{img_path})\n\n"
            md += f"*{cap_vi if vi else cap_en}*\n\n"
            md += md_table(["Field Name", "Description"], rows_for(lang, key))
            md += "\n"
    md += f"""## 4. Non-Functional Requirements

### 4.1 External Interfaces

- REST API prefix `/api/[resource]`; standard response `{{ "status": number, "message": string, "data": object }}`.
- MySQL 8 with utf8mb4 and UTC timestamps.
- File storage in `/uploads` or S3-compatible storage; no BLOB for media.
- SMTP for verification/reset/notification email with retry.
- OCR/Speech AI service invoked asynchronously with timeout, retry and fallback.

### 4.2 Quality Attributes

#### 4.2.1 Usability

Responsive web UI, Vietnamese-first learning experience, clear loading/error/empty states, and role-specific navigation.

#### 4.2.2 Performance

Common API responses should average under 2 seconds in normal conditions. AI requests return `job_id` immediately and process in the background.

#### 4.2.3 Security

JWT is mandatory for private APIs; bcrypt cost >= 10; backend checks Role + Subscription/Level; DTOs are mandatory for API responses; secrets are not stored in source control.

#### 4.2.4 Reliability, Auditability and Maintainability

Soft delete is used for important data, submitted attempts are immutable, important Staff/Admin operations are audited, global exception handling returns standard JSON, and business logic stays in backend Services.

## 5. Requirement Appendix

### 5.1 Business Rules

{md_table(["ID", "Rule Definition"], brs)}

### 5.2 System Messages

{md_table(["#", "Message code", "Message Type", "Context", "Content"], messages)}

### 5.3 Other Requirements

- No hard delete for important business data.
- No frontend scoring, authorization or subscription business logic.
- No schema change without migration.
- No hardcoded secrets/password/API keys.
- Source references: `AGENTS.md`, `CLAUDE.md`, `docs/01-SRS-Requirements/shared_context.md`, `docs/01-SRS-Requirements/use-cases/Bao_cao_dac_ta_Use_Case.md`.
"""
    return md


def make_docx(lang, path):
    vi = lang == "vi"
    doc = Document(str(TEMPLATE))
    clear_document(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    actors = base.actors_vi if vi else base.actors_en
    ucs = base.use_cases_vi if vi else base.use_cases_en
    brs = base.business_rules_vi if vi else base.business_rules_en
    entities = base.entities_vi if vi else base.entities_en
    non_ui = base.non_ui_vi if vi else base.non_ui_en
    messages = base.messages_vi if vi else base.messages_en
    screens = screen_descriptions_vi if vi else screen_descriptions_en
    specs = uc_specs_vi if vi else uc_specs_en
    title = "ĐẶC TẢ YÊU CẦU PHẦN MỀM" if vi else "SOFTWARE REQUIREMENT SPECIFICATION"

    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(str(LOGO), width=Inches(1.4))
    add_p(doc, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(18)
    add_p(doc, "Japanese Skill Practice Platform", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(16)
    add_p(doc, "- Hanoi, July 2026 -", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "Table of Contents", bold=True)
    add_toc_field(doc)

    add_h(doc, "I. Record of Changes", 1)
    add_table(doc, ["Date", "A/M/D", "In charge", "Change Description"], [
        ("2026-07-23", "A", "AI Agent", "Generated bilingual SRS according to Template1 DOCX structure."),
        ("2026-07-23", "M", "AI Agent", "Added missing Template1 content: screen descriptions, UC specifications, functional fields, NFRs and appendix."),
        ("2026-07-24", "M", "AI Agent", "Real Word Table of Contents field; split 1.2 into 6 sub-processes with swim-lane diagrams; added screen images to section 3 (real screenshots for Auth screens, wireframe mockups elsewhere)."),
        ("2026-07-24", "M", "AI Agent", "Expanded Use Case Specifications (section 2): longer step-by-step Normal Flow and named Alternative Flows for all 15 UC groups, written in simple, easy-to-understand English."),
        ("2026-07-24", "M", "AI Agent", "Split 1.3.3 Use Case Diagrams into one diagram per actor (Guest, Student, Staff, StaffManager, Admin), matching Template1's per-actor layout; added Guest to the Actors table."),
    ])
    add_p(doc, "*A - Added, M - Modified, D - Deleted")

    add_h(doc, "II. Software Requirement Specification", 1)
    add_h(doc, "1. Overall Requirements", 2)
    add_h(doc, "1.1 Context Diagram", 3)
    add_p(doc, "Japanese Skill Practice Platform " + ("là hệ thống web học và luyện thi tiếng Nhật JLPT từ N5 đến N1. Hệ thống có React frontend, Spring Boot REST API, MySQL database, file storage, SMTP và AI OCR/Speech service." if vi else "is a web-based JLPT learning and exam practice system from N5 to N1. It includes React frontend, Spring Boot REST API, MySQL database, file storage, SMTP, and OCR/Speech AI service."))
    add_picture_if_exists(doc, DIAGRAMS / "package-diagram.png")

    add_h(doc, "1.2 Main Business Processes", 3)
    process_sections = [
        ("1.2.1", "Authentication and Account Lifecycle", "Register, verify email, login, reset/change password, logout, token/session management.", "bp-1-authentication.png"),
        ("1.2.2", "Student Learning Path", "Browse level content, unlock lessons by order, complete lessons, bookmark and track progress.", "bp-2-learning-path.png"),
        ("1.2.3", "Assessment and SRS Review", "Take quizzes/mock exams, backend grading, immutable attempts, flashcard spaced repetition.", "bp-3-assessment-srs.png"),
        ("1.2.4", "AI Practice Processing", "Upload image/audio, create asynchronous AI jobs, timeout/retry/fallback, poll result.", "bp-4-ai-processing.png"),
        ("1.2.5", "Content Governance", "Staff drafts content, StaffManager reviews and publishes, audit and notification.", "bp-5-content-governance.png"),
        ("1.2.6", "System Administration", "Manage users, settings, subscriptions, reports, and notification rules.", "bp-6-system-admin.png"),
    ]
    for pid, name, desc, diagram in process_sections:
        add_h(doc, f"{pid} {name}", 4)
        add_p(doc, desc)
        add_picture_if_exists(doc, DIAGRAMS / diagram)

    add_h(doc, "1.3 User Requirements", 3)
    add_h(doc, "1.3.1 Actors", 4)
    add_table(doc, ["#", "Actor", "Description"], actors)
    add_h(doc, "1.3.2 Use Cases (UC)", 4)
    add_table(doc, ["ID", "Use Case", "Feature", "Use Case Description"], ucs)
    add_h(doc, "1.3.3 Use Case Diagrams", 4)
    add_h(doc, "1.3.3.1 UCs for Guest", 5)
    add_picture_if_exists(doc, DIAGRAMS / "uc-guest.png")
    add_h(doc, "1.3.3.2 UCs for Student", 5)
    add_picture_if_exists(doc, DIAGRAMS / "uc-student.png")
    add_h(doc, "1.3.3.3 UCs for Staff", 5)
    add_picture_if_exists(doc, DIAGRAMS / "uc-staff.png")
    add_h(doc, "1.3.3.4 UCs for StaffManager", 5)
    add_picture_if_exists(doc, DIAGRAMS / "uc-staffmanager.png")
    add_h(doc, "1.3.3.5 UCs for Admin", 5)
    add_picture_if_exists(doc, DIAGRAMS / "uc-admin.png")

    add_h(doc, "1.4 System Functionalities", 3)
    add_h(doc, "1.4.1 Screens Flow", 4)
    add_picture_if_exists(doc, DIAGRAMS / "screens-flow.png")
    add_p(doc, "Screen Descriptions", bold=True)
    add_table(doc, ["#", "Feature", "Screen (Route)", "Description"], screens)
    add_h(doc, "1.4.2 Screen Authorization", 4)
    add_table(doc, ["Screen", "Student", "Staff", "StaffManager", "Admin", "Notes"], base.screen_auth)
    add_h(doc, "1.4.3 Non-UI Functions", 4)
    add_table(doc, ["#", "Feature", "System Function", "Description"], non_ui)

    add_h(doc, "1.5 Entity Relationship Diagram", 3)
    add_picture_if_exists(doc, DIAGRAMS / "er-diagram.png")
    add_p(doc, "Entities Description", bold=True)
    add_table(doc, ["#", "Entity", "Description"], entities)

    add_h(doc, "2. Use Case Specifications", 2)
    current = None
    for spec in specs:
        if spec[0] != current:
            add_h(doc, spec[0], 3)
            current = spec[0]
        add_h(doc, spec[1], 4)
        add_uc_table(doc, spec)

    add_h(doc, "3. Functional Requirements", 2)
    groups = [
        ("3.1 User Authentication", ["User Register", "User Login", "Password Reset"]),
        ("3.2 Student Learning and Assessment", ["Lesson Detail", "Quiz / Mock Test Attempt", "AI Practice"]),
        ("3.3 Staff, Manager and Admin Operations", ["Staff Content Management", "Manager Review Queue", "Admin User Management"]),
    ]
    for group, keys in groups:
        add_h(doc, group, 3)
        for idx, key in enumerate(keys, 1):
            add_h(doc, f"{group.split()[0]}.{idx} {key}", 4)
            img_path, cap_en, cap_vi = screen_images[key]
            add_picture_if_exists(doc, DIAGRAMS / img_path)
            p = add_p(doc, cap_vi if vi else cap_en)
            p.runs[0].italic = True
            add_p(doc, "Field Description", bold=True)
            add_table(doc, ["Field Name", "Description"], rows_for(lang, key))

    add_h(doc, "4. Non-Functional Requirements", 2)
    add_h(doc, "4.1 External Interfaces", 3)
    add_table(doc, ["Interface", "Requirement"], [
        ("REST API", "/api/[resource], standard JSON {status,message,data}."),
        ("Database", "MySQL 8, utf8mb4, UTC timestamps, migration-controlled schema."),
        ("File Storage", "/uploads or S3-compatible storage; no media BLOB in DB."),
        ("Email", "SMTP with asynchronous sending and retry/outbox."),
        ("AI", "OCR/Speech services invoked asynchronously with timeout, retry and fallback."),
    ])
    add_h(doc, "4.2 Quality Attributes", 3)
    add_table(doc, ["Attribute", "Requirement"], [
        ("Usability", "Responsive Vietnamese-first UI with clear loading/error/empty states and role-specific navigation."),
        ("Performance", "Common API average response under 2 seconds; AI returns job_id immediately and processes in background."),
        ("Security", "JWT for private APIs, bcrypt cost >= 10, DTO-only responses, Role + Subscription/Level checks, no secrets in source."),
        ("Reliability", "Soft delete, immutable submitted attempts, retry/fallback for AI/email, global exception handler."),
        ("Maintainability", "Controller -> Service -> Repository -> Entity; business logic remains in backend Services."),
    ])

    add_h(doc, "5. Requirement Appendix", 2)
    add_h(doc, "5.1 Business Rules", 3)
    add_table(doc, ["ID", "Rule Definition"], brs)
    add_h(doc, "5.2 System Messages", 3)
    add_table(doc, ["#", "Message code", "Message Type", "Context", "Content"], messages)
    add_h(doc, "5.3 Other Requirements", 3)
    add_table(doc, ["#", "Requirement"], [
        ("1", "No hard delete for important business data."),
        ("2", "No frontend scoring, authorization or subscription business logic."),
        ("3", "No schema change without migration."),
        ("4", "No hardcoded secrets/password/API keys."),
        ("5", "Source references: AGENTS.md, CLAUDE.md, shared_context.md, Bao_cao_dac_ta_Use_Case.md."),
    ])
    doc.save(path)


def main():
    (OUT / "SRS-Japanese-Skill-Practice-Platform.vi.md").write_text(make_markdown("vi"), encoding="utf-8", newline="\n")
    (OUT / "SRS-Japanese-Skill-Practice-Platform.en.md").write_text(make_markdown("en"), encoding="utf-8", newline="\n")
    make_docx("vi", OUT / "SRS-Japanese-Skill-Practice-Platform.vi.docx")
    make_docx("en", OUT / "SRS-Japanese-Skill-Practice-Platform.en.docx")


if __name__ == "__main__":
    main()
